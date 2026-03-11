from pathlib import Path
import docx

in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新.docx")
out_path = in_path.with_name(in_path.stem + '_精简技术细节.docx')

doc = docx.Document(str(in_path))
body = doc._body._body

# locate 6.1
start_idx = None
for p in doc.paragraphs:
    if p.text.strip() == '6.1 工作总结':
        start_idx = body.index(p._element)
        break
if start_idx is None:
    raise SystemExit('6.1 not found')

# remove all after 6.1
for i in range(len(body)-1, start_idx, -1):
    body.remove(body[i])

insert_idx = start_idx + 1

lines = [
    '本文完成了基于Spring Boot的多人在线会议系统设计与实现。后端采用Spring Boot + MyBatis，MySQL存储业务数据，Redis缓存会话与热点数据，MinIO承载头像与文件；前端基于Vue 3，接口按RESTful规范统一前缀，关键链路使用JWT鉴权，结合WebSocket承载实时消息。',
    '系统实现了用户注册登录、信息与设置管理、联系人与申请、会议创建/加入/结束、预约会议、会议成员管理、通知中心等核心功能；会议内实时音视频通过WebRTC建立P2P连接，配合信令与状态同步完成入会、退会与控制；AI接口覆盖对话、摘要与建议，提升会议产出效率。测试结果表明主要业务流程稳定，满足中小规模会议场景。',
    '6.2 存在的不足',
    '1. 多方会议采用Mesh拓扑，规模扩大时客户端负载上升明显。',
    '2. 弱网下自适应策略与重传机制仍有提升空间。',
    '3. 移动端体验依赖浏览器，功能与性能受限。',
    '4. 录制、回放与会议资料沉淀能力不足。',
    '5. 监控、告警与运维体系尚不完善。',
    '6.3 未来展望',
    '1. 引入SFU转发架构，提升大规模会议稳定性。',
    '2. 优化码率控制、FEC/重传策略，增强弱网适配。',
    '3. 增强移动端能力与交互体验。',
    '4. 增加录制与回放能力，支持会议资料管理。',
    '5. 基于现有AI接口扩展纪要、要点提取与行动项跟踪。',
    '6. 建立监控与日志追踪体系，提升可维护性。',
]

for text in lines:
    p = doc.add_paragraph(text)
    body.remove(p._element)
    body.insert(insert_idx, p._element)
    insert_idx += 1

_doc = doc
_doc.save(str(out_path))
print('updated', out_path)
