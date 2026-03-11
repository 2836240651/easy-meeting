from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v6.docx")
OUT = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")

TEXT_STYLES = {"First Paragraph", "Normal", "Body Text"}


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def style_body_item(paragraph):
    paragraph.style = "Body Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run)


doc = Document(str(SRC))
i = 0
while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    if para.style.name == "Heading 4":
        next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
        title = para.text.strip()
        if next_para is not None and next_para.style.name in TEXT_STYLES and next_para.text.strip():
            body = next_para.text.strip().lstrip("：:")
            if "：" not in title and ":" not in title:
                para.text = f"{title}：{body}"
            else:
                para.text = title
            style_body_item(para)
            delete_paragraph(next_para)
            continue
        style_body_item(para)
    i += 1

doc.save(str(OUT))
print(OUT)
