from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_PATH = Path(r"C:\Users\13377\Desktop\论文相关\第2章_相关技术介绍.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_format_backup.docx")
OUTPUT_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_已调整格式.docx")


def set_east_asia_font(run, font_name: str, size_pt: float, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def append_page_number(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)
    set_east_asia_font(run, "宋体", 9, False)


def ensure_rpr(style):
    style_el = style._element
    rpr = style_el.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style_el.append(rpr)
    return rpr


def set_style_font(style, ascii_font: str, east_asia_font: str, size_pt: float, bold: bool = False):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    font.bold = bold
    rpr = ensure_rpr(style)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def format_paragraph(paragraph, style_name: str):
    fmt = paragraph.paragraph_format

    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.space_before = Pt(18)
        fmt.space_after = Pt(18)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif style_name == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif style_name == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(3)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif style_name == "Source Code":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.left_indent = Cm(0.74)
        fmt.right_indent = Cm(0.74)
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(3)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.first_line_indent = Cm(0.74)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


if not BACKUP_PATH.exists():
    shutil.copy2(DOC_PATH, BACKUP_PATH)

doc = Document(str(DOC_PATH))

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

styles = doc.styles
set_style_font(styles["Normal"], "Times New Roman", "宋体", 12)
set_style_font(styles["Body Text"], "Times New Roman", "宋体", 12)
set_style_font(styles["First Paragraph"], "Times New Roman", "宋体", 12)
set_style_font(styles["Heading 1"], "Times New Roman", "黑体", 15, True)
set_style_font(styles["Heading 2"], "Times New Roman", "黑体", 14, True)
set_style_font(styles["Heading 3"], "Times New Roman", "黑体", 12, True)
set_style_font(styles["Source Code"], "Consolas", "等线", 10.5)

for style_name in ("Normal", "Body Text", "First Paragraph"):
    pf = styles[style_name].paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

for style_name in ("Heading 2", "Heading 3"):
    pf = styles[style_name].paragraph_format
    pf.line_spacing = 1.5

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if not text:
        continue

    style_name = paragraph.style.name
    if text == "参考文献":
        paragraph.style = styles["Heading 1"]
        style_name = "Heading 1"
    elif text == "本章小结":
        paragraph.style = styles["Heading 2"]
        style_name = "Heading 2"
    elif text.startswith("[") and style_name in {"Body Text", "First Paragraph", "Normal"}:
        paragraph.style = styles["Body Text"]
        style_name = "Body Text"
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.hanging_indent = Cm(0.74)

    format_paragraph(paragraph, style_name)

    for run in paragraph.runs:
        if style_name == "Source Code":
            set_east_asia_font(run, "Consolas", 10.5, False)
        elif style_name == "Heading 1":
            set_east_asia_font(run, "黑体", 15, True)
        elif style_name == "Heading 2":
            set_east_asia_font(run, "黑体", 14, True)
        elif style_name == "Heading 3":
            set_east_asia_font(run, "黑体", 12, True)
        else:
            set_east_asia_font(run, "宋体", 12, False)
            if any("A" <= ch <= "z" for ch in run.text):
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

footer = section.footer
footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_p.clear()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_fmt = footer_p.paragraph_format
footer_fmt.first_line_indent = None
footer_fmt.left_indent = None
footer_fmt.right_indent = None
footer_fmt.space_before = Pt(0)
footer_fmt.space_after = Pt(0)
footer_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
append_page_number(footer_p)

doc.save(str(OUTPUT_PATH))
print(f"Formatted: {OUTPUT_PATH}")
print(f"Backup: {BACKUP_PATH}")
