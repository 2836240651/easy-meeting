import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SRC = r"C:\Users\13377\Desktop\论文相关\1.0.1.docx"
BASE, EXT = os.path.splitext(SRC)
OUT = BASE + "_v1.0.2" + EXT

TITLE_SIZE = Pt(10.5)  # 五号
BODY_SIZE = Pt(9)      # 小五
BLACK = RGBColor(0, 0, 0)


def is_cjk(ch: str) -> bool:
    code = ord(ch)
    return (
        0x4E00 <= code <= 0x9FFF
        or 0x3400 <= code <= 0x4DBF
        or 0x3000 <= code <= 0x303F
        or 0xFF00 <= code <= 0xFFEF
    )


def split_segments(text: str):
    if not text:
        return []
    segments = []
    current = text[0]
    current_kind = is_cjk(text[0])
    for ch in text[1:]:
        kind = is_cjk(ch)
        if kind == current_kind:
            current += ch
        else:
            segments.append((current, current_kind))
            current = ch
            current_kind = kind
    segments.append((current, current_kind))
    return segments


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_run_font(run, chinese: bool, size, bold: bool):
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    if chinese:
        run.font.name = "宋体"
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
    else:
        run.font.name = "Times New Roman"
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")


def rewrite_paragraph(paragraph, size, bold, center=False):
    text = paragraph.text
    clear_paragraph(paragraph)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for seg, chinese in split_segments(text):
        run = paragraph.add_run(seg)
        set_run_font(run, chinese, size, bold)
    for run in paragraph.runs:
        run.underline = False
        run.italic = False


def insert_blank_paragraph_before(table):
    tbl = table._element
    blank_p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    blank_p.append(p_pr)
    tbl.addprevious(blank_p)


def iter_block_items(document):
    body = document._element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", child)
        elif child.tag == qn("w:tbl"):
            yield ("t", child)


def find_paragraph_by_element(document, element):
    for paragraph in document.paragraphs:
        if paragraph._element == element:
            return paragraph
    return None


def find_table_by_element(document, element):
    for table in document.tables:
        if table._element == element:
            return table
    return None


def is_table_title(text: str) -> bool:
    stripped = text.strip()
    return bool(stripped) and (stripped.startswith("表") or stripped.startswith("续表"))


doc = Document(SRC)

# 1. 表格正文格式
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    rewrite_paragraph(paragraph, BODY_SIZE, bold=False, center=False)
                else:
                    for run in paragraph.runs:
                        set_run_font(run, True, BODY_SIZE, False)

# 2. 表格标题格式 + 标题与表格空一行
blocks = list(iter_block_items(doc))
for idx, (kind, element) in enumerate(blocks):
    if kind != "t":
        continue
    table = find_table_by_element(doc, element)
    if table is None:
        continue

    prev_idx = idx - 1
    while prev_idx >= 0 and blocks[prev_idx][0] == "p":
        prev_para = find_paragraph_by_element(doc, blocks[prev_idx][1])
        if prev_para is None:
            break
        text = prev_para.text.strip()
        if not text:
            prev_idx -= 1
            continue
        if is_table_title(text):
            rewrite_paragraph(prev_para, TITLE_SIZE, bold=True, center=True)
            if idx - prev_idx == 1:
                insert_blank_paragraph_before(table)
        break

doc.save(OUT)
print(OUT)
