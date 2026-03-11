from docx import Document


path = r"C:\Users\13377\Desktop\论文相关\1.0.1.docx"
doc = Document(path)

print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))

for i, p in enumerate(doc.paragraphs[:160]):
    txt = p.text.strip().replace("\n", " ")
    if txt:
        print(f"P{i}: {txt[:120]}")

for ti, table in enumerate(doc.tables[:20], start=1):
    print(f"Table {ti}: {len(table.rows)}x{len(table.columns)}")
    for r in range(min(len(table.rows), 3)):
        vals = [cell.text.strip().replace("\n", " ") for cell in table.rows[r].cells]
        print("  ", vals[:8])
