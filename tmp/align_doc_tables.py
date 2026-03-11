import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SRC = r"C:\Users\13377\Desktop\论文相关\1.0.1_v1.0.2.docx"
BASE, EXT = os.path.splitext(SRC)
OUT = BASE[:-7] + "_v1.0.3" + EXT if BASE.endswith("_v1.0.2") else BASE + "_v1.0.3" + EXT

doc = Document(SRC)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                fmt = paragraph.paragraph_format
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt.first_line_indent = Pt(0)
                fmt.left_indent = Pt(0)
                fmt.right_indent = Pt(0)

doc.save(OUT)
print(OUT)
