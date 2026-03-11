import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SRC = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节.docx")
DST = SRC.with_name(SRC.stem + "_第5章表题字号字体修正_v3.docx")

CHAPTER = 5
SONG = "宋体"
TNR = "Times New Roman"
SIZE = Pt(10.5)  # 五号

PUNCT = set("，。；：？！、,.!?;:（）()【】[]《》<>“”‘’·/\\|—–－")

SECTION_RE = re.compile(r"^5\.(\d+)(?:\.(\d+))?\s+(.+?)\s*$")
CAPTION_WITH_TABLE_RE = re.compile(r"^(续表)?\s*表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")
CAPTION_CONT_RE = re.compile(r"^续表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def is_chapter_heading(text: str, chapter: int) -> bool:
    t = text.strip()
    cn = "一二三四五六七八九十"[chapter-1]
    return t.startswith(f"第{chapter}章") or t.startswith(f"第{cn}章")


def set_run_fonts(run, east_asia: str, ascii_font: str, size=SIZE, bold=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    run.font.name = ascii_font


def sanitize_title(title: str) -> str:
    t = (title or "").strip()
    t = re.sub(r"^[0-9]+(?:\.[0-9]+)+\s*", "", t)
    t = re.sub(r"^5\.[0-9]+(?:\.[0-9]+)?\s*", "", t)
    t = re.sub(r"^（\d+）\s*", "", t)
    t = re.sub(r"^\(\d+\)\s*", "", t)
    t = re.sub(r"^\d+[\.)、]\s*", "", t)
    t = re.sub(r"[：:]\s*$", "", t)
    t = "".join(ch for ch in t if ch not in PUNCT)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def insert_paragraph_before(tbl: Table) -> Paragraph:
    p = OxmlElement('w:p')
    tbl._element.addprevious(p)
    return Paragraph(p, tbl._parent)


def make_caption(paragraph: Paragraph, label: str, title: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)

    r1 = paragraph.add_run("表")
    set_run_fonts(r1, SONG, TNR, bold=True)
    r2 = paragraph.add_run(label)
    set_run_fonts(r2, TNR, TNR, bold=True)
    r3 = paragraph.add_run(" " + title)
    set_run_fonts(r3, SONG, TNR, bold=True)

    pPr = paragraph._element.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def format_table_fonts(tbl: Table):
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

    if tbl.rows:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))

    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_fonts(run, SONG, TNR, size=SIZE, bold=None)


def extract_old_label_and_title(text: str):
    tx = (text or "").strip()
    if not tx:
        return None, "", False

    m = CAPTION_WITH_TABLE_RE.match(tx)
    if m:
        is_cont = bool(m.group(1))
        return m.group(2), sanitize_title(m.group(3)), is_cont

    m = CAPTION_CONT_RE.match(tx)
    if m:
        return m.group(1), sanitize_title(m.group(2)), True

    return None, "", False


def is_testcase_table(header: str) -> bool:
    return (
        '用例编号' in header or
        ('测试步骤' in header and '预期结果' in header) or
        ('测试方法' in header and '预期结果' in header) or
        ('测试场景' in header and '预期结果' in header)
    )


def choose_title(section_title: str, prev_text: str, header: str, had_caption: bool, caption_title: str, is_cont: bool) -> str:
    if had_caption and caption_title:
        return caption_title

    base = sanitize_title(section_title) or ""

    # Strong signals from table header
    if '接口路径' in header and '方法' in header:
        return sanitize_title("RESTful API接口测试续表" if is_cont else "RESTful API接口测试结果表")

    if is_testcase_table(header):
        if base:
            suffix = "用例表" if base.endswith("测试") else "测试用例表"
            return sanitize_title(base + suffix)
        return "测试用例表"

    if ('TPS' in header) or ('平均响应时间' in header) or ('并发数' in header and '平均响应时间' in header):
        return sanitize_title((base or "接口性能测试") + "结果表")

    if ('桌面应用' in header) or ('Web端' in header) or ('浏览器' in header and '测试结果' in header):
        return sanitize_title((base or "兼容性测试") + "结果表")

    if ('带宽' in header) or ('音视频延迟' in header) or ('画面质量' in header):
        return sanitize_title((base or "网络环境测试") + "结果表")

    if ('配置说明' in header) or ('设备类型' in header) or ('操作系统' in header and '浏览器' in header):
        if '客户端' in prev_text:
            return sanitize_title("客户端测试环境配置表")
        return sanitize_title("测试环境配置表")

    # If it's a continuation caption but title couldn't be parsed
    if is_cont and not caption_title:
        return sanitize_title("RESTful API接口测试续表")

    # Fallback to previous paragraph if it's short and not sentence-like
    prev_candidate = sanitize_title(prev_text)
    if prev_candidate and (len(prev_candidate) <= 14) and not any(w in prev_candidate for w in ["以下", "进行", "如下", "本节", "主要", "包括"]):
        if prev_candidate not in {"测试用例", "测试场景", "测试结果", "性能测试结果", "安全测试用例", "主要接口测试结果", "压力测试结果"}:
            if not prev_candidate.endswith('表'):
                prev_candidate += '表'
            return sanitize_title(prev_candidate)

    if base:
        if not base.endswith('表'):
            base += '表'
        return sanitize_title(base)

    return "表"


def update_single_paragraph_text(p: Paragraph, new_text: str):
    for r in list(p.runs):
        p._element.remove(r._element)
    r = p.add_run(new_text)
    set_run_fonts(r, SONG, TNR, size=SIZE, bold=None)


doc = Document(SRC)
blocks = list(iter_block_items(doc))

in_ch5 = False
section_title = ""
records = []
ch5_paragraphs = []

for idx, b in enumerate(blocks):
    if isinstance(b, Paragraph):
        t = (b.text or "").strip()
        if is_chapter_heading(t, CHAPTER):
            in_ch5 = True
        elif is_chapter_heading(t, CHAPTER + 1) and in_ch5:
            in_ch5 = False
        if in_ch5:
            ch5_paragraphs.append(b)
            m = SECTION_RE.match(t)
            if m:
                section_title = m.group(3).strip()
    elif isinstance(b, Table) and in_ch5:
        prev_para = None
        prev_text = ""
        j = idx - 1
        while j >= 0:
            if isinstance(blocks[j], Paragraph):
                tx = (blocks[j].text or "").strip()
                if tx:
                    prev_para = blocks[j]
                    prev_text = tx
                    break
            elif isinstance(blocks[j], Table):
                break
            j -= 1

        header = ""
        if b.rows:
            header = " ".join(c.text.strip().replace('\n',' ') for c in b.rows[0].cells)

        had_caption = prev_text.startswith('表') or prev_text.startswith('续表')
        old_label, cap_title, is_cont = extract_old_label_and_title(prev_text) if had_caption else (None, "", False)

        if is_cont and cap_title and '续表' not in cap_title:
            cap_title = sanitize_title(cap_title + "续表")

        records.append({
            'table': b,
            'prev_para': prev_para,
            'prev_text': prev_text,
            'section_title': section_title,
            'header': header,
            'had_caption': had_caption,
            'old_label': old_label,
            'caption_title': cap_title,
            'is_cont': is_cont,
        })

print(f"Chapter {CHAPTER} tables: {len(records)}")

old_to_new = {}
caption_paras = set()

for seq, rec in enumerate(records, start=1):
    new_label = f"{CHAPTER}-{seq}"
    if rec['old_label']:
        old_to_new[rec['old_label']] = new_label

    title = choose_title(
        section_title=rec['section_title'],
        prev_text=rec['prev_text'],
        header=rec['header'],
        had_caption=rec['had_caption'],
        caption_title=rec['caption_title'],
        is_cont=rec['is_cont'],
    )
    title = sanitize_title(title) or "表"

    if rec['had_caption'] and rec['prev_para'] is not None:
        caption_para = rec['prev_para']
    else:
        caption_para = insert_paragraph_before(rec['table'])

    make_caption(caption_para, new_label, title)
    caption_paras.add(caption_para._element)

    format_table_fonts(rec['table'])

# Update in-text references (skip caption paras)
for p in ch5_paragraphs:
    if p._element in caption_paras:
        continue
    text = p.text or ""
    new = text
    for old, new_label in old_to_new.items():
        new = re.sub(rf"表\s*[-–—－]?\s*{re.escape(old)}(?!\d)", f"表{new_label}", new)
    if new != text:
        update_single_paragraph_text(p, new)


doc.save(DST)
print(f"Saved: {DST}")
print(f"Old->New mapping: {old_to_new}")
