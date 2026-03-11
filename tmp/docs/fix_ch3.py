import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

INP = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终.docx")
OUT = INP.with_name(INP.stem + "_第3章同步修正.docx")

CHAPTER = 3
SONG = "宋体"
TNR = "Times New Roman"
SIZE = Pt(10.5)  # 五号

# include '_' as punctuation for table names
PUNCT = set("，。；：？！、,.!?;:（）()【】[]《》<>“”‘’·/\\|—–－_-\t")

CAPTION_WITH_TABLE_RE = re.compile(r"^(续表)?\s*表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")
CAPTION_CONT_RE = re.compile(r"^续表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def is_chapter_heading(text: str, chapter: int) -> bool:
    t = text.strip()
    cn = "一二三四五六七八九十"[chapter-1]
    return t.startswith(f"第{chapter}章") or t.startswith(f"第{cn}章")


def set_run_fonts(run, east_asia: str, ascii_font: str, size=SIZE, bold=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    run.font.name = ascii_font


def sanitize_title(title: str) -> str:
    t = (title or "").strip()
    # remove common leading numbering/colons
    t = re.sub(r"^[0-9]+(?:\.[0-9]+)+\s*", "", t)
    t = re.sub(r"^（\d+）\s*", "", t)
    t = re.sub(r"^\(\d+\)\s*", "", t)
    t = re.sub(r"^\d+[\.)、]\s*", "", t)
    t = re.sub(r"[：:]\s*$", "", t)

    # convert underscores to spaces before stripping punctuation
    t = t.replace('_', ' ')
    t = "".join(ch for ch in t if ch not in PUNCT)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def insert_paragraph_before(tbl: Table) -> Paragraph:
    p = OxmlElement('w:p')
    tbl._element.addprevious(p)
    return Paragraph(p, tbl._parent)


def make_caption(paragraph: Paragraph, label: str, title: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)

    r1 = paragraph.add_run("表")
    set_run_fonts(r1, SONG, TNR, bold=True)
    r2 = paragraph.add_run(label)
    set_run_fonts(r2, TNR, TNR, bold=True)
    r3 = paragraph.add_run(" " + title)
    set_run_fonts(r3, SONG, TNR, bold=True)

    pPr = paragraph._element.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def format_table_fonts(tbl: Table):
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

    if tbl.rows:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))

    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_fonts(run, SONG, TNR, size=SIZE, bold=None)


def extract_old_label_and_title(text: str):
    tx = (text or "").strip()
    if not tx:
        return None, "", False

    m = CAPTION_WITH_TABLE_RE.match(tx)
    if m:
        is_cont = bool(m.group(1))
        return m.group(2), sanitize_title(m.group(3)), is_cont

    m = CAPTION_CONT_RE.match(tx)
    if m:
        return m.group(1), sanitize_title(m.group(2)), True

    return None, "", False


def is_blank_table(tbl: Table) -> bool:
    for row in tbl.rows:
        for cell in row.cells:
            if cell.text and cell.text.strip():
                return False
    return True


def update_single_paragraph_text(p: Paragraph, new_text: str):
    for r in list(p.runs):
        p._element.remove(r._element)
    r = p.add_run(new_text)
    set_run_fonts(r, SONG, TNR, size=SIZE, bold=None)


doc = Document(INP)
blocks = list(iter_block_items(doc))

in_ch = False
records = []
ch_paragraphs = []

for idx, b in enumerate(blocks):
    if isinstance(b, Paragraph):
        t = (b.text or "").strip()
        if is_chapter_heading(t, CHAPTER):
            in_ch = True
        elif is_chapter_heading(t, CHAPTER + 1) and in_ch:
            in_ch = False
        if in_ch:
            ch_paragraphs.append(b)
    elif isinstance(b, Table) and in_ch:
        if is_blank_table(b):
            continue
        # immediate prev non-empty paragraph
        prev_para = None
        prev_text = ""
        j = idx - 1
        while j >= 0:
            if isinstance(blocks[j], Paragraph):
                tx = (blocks[j].text or "").strip()
                if tx:
                    prev_para = blocks[j]
                    prev_text = tx
                    break
            elif isinstance(blocks[j], Table):
                break
            j -= 1

        had_caption = prev_text.startswith('表') or prev_text.startswith('续表')
        old_label, cap_title, is_cont = extract_old_label_and_title(prev_text) if had_caption else (None, "", False)

        records.append({
            'table': b,
            'prev_para': prev_para,
            'prev_text': prev_text,
            'had_caption': had_caption,
            'old_label': old_label,
            'caption_title': cap_title,
            'is_cont': is_cont,
        })

print(f"Chapter {CHAPTER} non-blank tables: {len(records)}")

old_to_new = {}
caption_paras = set()

for seq, rec in enumerate(records, start=1):
    new_label = f"{CHAPTER}-{seq}"
    if rec['old_label']:
        old_to_new[rec['old_label']] = new_label

    title = rec['caption_title']
    if not title:
        # fallback to previous paragraph text
        title = sanitize_title(rec['prev_text'])

    title = sanitize_title(title)
    if title and not title.endswith('表'):
        title += '表'
    title = title or '表'

    if rec['had_caption'] and rec['prev_para'] is not None:
        cap_para = rec['prev_para']
    else:
        cap_para = insert_paragraph_before(rec['table'])

    make_caption(cap_para, new_label, title)
    caption_paras.add(cap_para._element)

    format_table_fonts(rec['table'])

# Update in-text refs inside chapter (skip caption paras)
for p in ch_paragraphs:
    if p._element in caption_paras:
        continue
    text = p.text or ""
    new = text
    for old, new_label in old_to_new.items():
        new = re.sub(rf"表\s*[-–—－]?\s*{re.escape(old)}(?!\d)", f"表{new_label}", new)
    if new != text:
        update_single_paragraph_text(p, new)


doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Old->New mapping: {old_to_new}")
