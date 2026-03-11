import re
from pathlib import Path
from docx import Document

INP = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正.docx")
OUT = INP.with_name(INP.stem + "_第4章代码按项目更新.docx")

ROOT = Path(r"D:\JavaPartical\easymeeting-java")

# Backend files
POM = ROOT / "pom.xml"
ACCOUNT_CTRL = ROOT / "src/main/java/com/easymeeting/controller/AccountController.java"
MEETING_INFO_CTRL = ROOT / "src/main/java/com/easymeeting/controller/MeetingInfoController.java"
MEETING_RESERVE_CTRL = ROOT / "src/main/java/com/easymeeting/controller/MeetingReserveController.java"
CHAT_CTRL = ROOT / "src/main/java/com/easymeeting/controller/ChatController.java"
USER_CONTACT_CTRL = ROOT / "src/main/java/com/easymeeting/controller/UserContactController.java"
NOTIFICATION_CTRL = ROOT / "src/main/java/com/easymeeting/controller/UserNotificationController.java"
APP_MAIN = ROOT / "src/main/java/com/easymeeting/EasymeetingApplication.java"

USER_INFO_MAPPER = ROOT / "src/main/java/com/easymeeting/mappers/UserInfoMapper.java"
MEETING_INFO_MAPPER = ROOT / "src/main/java/com/easymeeting/mappers/MeetingInfoMapper.java"
MEETING_INFO_SERVICE_IMPL = ROOT / "src/main/java/com/easymeeting/service/impl/MeetingInfoServiceImpl.java"

REDIS_CONFIG = ROOT / "src/main/java/com/easymeeting/redis/RedisConfig.java"
REDISSON_CONFIG = ROOT / "src/main/java/com/easymeeting/redis/RedissonConfig.java"
GLOBAL_INTERCEPTOR = ROOT / "src/main/java/com/easymeeting/annotation/globalInterceptor.java"
GLOBAL_ASPECT = ROOT / "src/main/java/com/easymeeting/aspect/globalOperationAspect.java"
GLOBAL_EXCEPTION = ROOT / "src/main/java/com/easymeeting/controller/AGlobalExceptionHandlerController.java"

NETTY_STARTER = ROOT / "src/main/java/com/easymeeting/websocket/netty/NettyWebSocketStarter.java"
NETTY_HANDLER = ROOT / "src/main/java/com/easymeeting/websocket/netty/HandlerWebSocket.java"
REDIS_COMPONENT = ROOT / "src/main/java/com/easymeeting/redis/RedisComponent.java"
MINIO_CONFIG = ROOT / "src/main/java/com/easymeeting/entity/config/MinioConfig.java"
MINIO_SERVICE_IMPL = ROOT / "src/main/java/com/easymeeting/service/impl/MinioServiceImpl.java"

# Frontend files
FRONTEND = ROOT / "frontend"
PKG = FRONTEND / "package.json"
WEBRTC = FRONTEND / "src/api/webrtc-manager.js"
ROUTER = FRONTEND / "src/router/index.js"
MEETING_VUE = FRONTEND / "src/views/Meeting.vue"
AXIOS_JS = FRONTEND / "src/api/axios.js"
SERVICES_JS = FRONTEND / "src/api/services.js"
LOGIN_VUE = FRONTEND / "src/views/Login.vue"
ELECTRON_MAIN = FRONTEND / "electron/main.js"
ELECTRON_PRELOAD = FRONTEND / "electron/preload.js"
SETTINGS_MANAGER = FRONTEND / "src/utils/settings-manager.js"


def read_text(path: Path) -> str:
    # tolerate BOM
    return path.read_text(encoding="utf-8-sig")


def find_chapter_range(doc, chapter: int):
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if start is None and (t.startswith(f"第{chapter}章") or t.startswith("第" + "一二三四五六七八九十"[chapter-1] + "章")):
            start = i
        elif start is not None and end is None and (
            t.startswith(f"第{chapter+1}章") or t.startswith("第" + "一二三四五六七八九十"[chapter] + "章")
        ):
            end = i
            break
    return start, end


def brace_slice(src: str, open_brace_pos: int) -> str:
    depth = 0
    i = open_brace_pos
    in_squote = False
    in_dquote = False
    esc = False
    in_line_comment = False
    in_block_comment = False

    while i < len(src):
        ch = src[i]
        nxt = src[i+1] if i+1 < len(src) else ""

        if in_line_comment:
            if ch == "\n":
                in_line_comment = False
            i += 1
            continue

        if in_block_comment:
            if ch == "*" and nxt == "/":
                in_block_comment = False
                i += 2
            else:
                i += 1
            continue

        if in_squote:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == "'":
                in_squote = False
            i += 1
            continue

        if in_dquote:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_dquote = False
            i += 1
            continue

        # not in string/comment
        if ch == "/" and nxt == "/":
            in_line_comment = True
            i += 2
            continue
        if ch == "/" and nxt == "*":
            in_block_comment = True
            i += 2
            continue
        if ch == "'":
            in_squote = True
            i += 1
            continue
        if ch == '"':
            in_dquote = True
            i += 1
            continue

        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return src[open_brace_pos: i+1]

        i += 1

    raise ValueError("Unbalanced braces")


def extract_java_method(path: Path, method_name: str) -> str:
    s = read_text(path)
    m = re.search(rf"\b{re.escape(method_name)}\s*\(", s)
    if not m:
        raise ValueError(f"Method {method_name} not found in {path}")

    # find the line containing the method signature (search backwards for 'public' in nearby window)
    sig_pos = m.start()
    window_start = max(0, sig_pos - 2000)
    pub = s.rfind("public", window_start, sig_pos)
    if pub == -1:
        pub = s.rfind("protected", window_start, sig_pos)
    if pub == -1:
        raise ValueError(f"Cannot locate signature for {method_name} in {path}")

    # expand to include preceding annotations
    start_line = s.rfind("\n", 0, pub) + 1
    scan = start_line
    while True:
        prev_nl = s.rfind("\n", 0, scan-1)
        if prev_nl == -1:
            break
        line = s[prev_nl+1:scan].rstrip("\r\n")
        stripped = line.strip()
        if stripped.startswith("@") or stripped.startswith("//") or stripped.startswith("/*") or stripped.startswith("*"):
            scan = prev_nl + 1
            continue
        # stop on blank line or non-annotation
        if stripped == "":
            break
        break
    start = scan

    # find first '{' after pub
    brace = s.find("{", pub)
    if brace == -1:
        raise ValueError("No brace")

    body = brace_slice(s, brace)
    # include from start to end of body
    snippet = s[start: brace] + body
    return snippet.strip("\n\r")


def extract_java_type(path: Path, type_name: str, kind: str = "class") -> str:
    s = read_text(path)
    m = re.search(rf"\bpublic\s+{kind}\s+{re.escape(type_name)}\b", s)
    if not m:
        # interfaces in this repo sometimes omit 'public' on BaseMapper
        m = re.search(rf"\b{kind}\s+{re.escape(type_name)}\b", s)
    if not m:
        raise ValueError(f"{kind} {type_name} not found in {path}")
    start = s.rfind("\n", 0, m.start()) + 1
    brace = s.find("{", m.end())
    body = brace_slice(s, brace)
    snippet = s[start:brace] + body
    return snippet.strip("\n\r")


def extract_between_tags(text: str, start_tag: str, end_tag: str) -> str:
    a = text.find(start_tag)
    if a == -1:
        raise ValueError(f"start tag {start_tag} not found")
    b = text.find(end_tag, a)
    if b == -1:
        raise ValueError(f"end tag {end_tag} not found")
    return text[a: b + len(end_tag)].strip("\n\r")


def extract_pom_dependencies() -> str:
    s = read_text(POM)
    m = re.search(r"<dependencies>", s)
    if not m:
        raise ValueError("<dependencies> not found")
    start = m.start()
    end = s.find("</dependencies>", start)
    if end == -1:
        raise ValueError("</dependencies> not found")
    end = end + len("</dependencies>")
    # keep indentation and a compact excerpt: full dependencies section is acceptable
    return s[start:end].strip()


def extract_package_dependencies() -> str:
    s = read_text(PKG)
    # small snippet with dependencies and devDependencies
    m = re.search(r'"dependencies"\s*:\s*\{', s)
    if not m:
        raise ValueError('dependencies not found')
    # naive brace match starting at '{'
    brace = s.find('{', m.end()-1)
    deps_body = brace_slice(s, brace)

    # include devDependencies if present
    md = re.search(r'"devDependencies"\s*:\s*\{', s)
    dev = ""
    if md:
        b2 = s.find('{', md.end()-1)
        dev_body = brace_slice(s, b2)
        dev = f"\n  \"devDependencies\": {dev_body}"

    snippet = "{\n  \"dependencies\": " + deps_body + dev + "\n}"
    return snippet.strip()


def extract_js_object_literal(path: Path, start_token: str) -> str:
    s = read_text(path)
    pos = s.find(start_token)
    if pos == -1:
        raise ValueError(f"token not found: {start_token}")
    # from start of line
    start = s.rfind("\n", 0, pos) + 1
    # find '{' after token
    brace = s.find("{", pos)
    if brace == -1:
        raise ValueError("no brace")
    body = brace_slice(s, brace)
    # take until end of body
    return (s[start:brace] + body).strip("\n\r")


def extract_js_method(path: Path, method_sig_regex: str) -> str:
    s = read_text(path)
    m = re.search(method_sig_regex, s)
    if not m:
        raise ValueError(f"method not found: {method_sig_regex}")
    start = s.rfind("\n", 0, m.start()) + 1

    # Find the function body brace '{' after the parameter list ')'.
    # This avoids accidentally selecting '{' inside default parameters like `options = {}`.
    open_paren = s.find("(", m.start(), m.end() + 5)
    if open_paren == -1:
        open_paren = s.find("(", m.start())
    if open_paren == -1:
        raise ValueError("No '(' found for method signature")

    # match parentheses while skipping strings/comments
    depth = 0
    i = open_paren
    in_squote = False
    in_dquote = False
    esc = False
    in_line_comment = False
    in_block_comment = False
    close_paren = None
    while i < len(s):
        ch = s[i]
        nxt = s[i + 1] if i + 1 < len(s) else ""

        if in_line_comment:
            if ch == "\n":
                in_line_comment = False
            i += 1
            continue
        if in_block_comment:
            if ch == "*" and nxt == "/":
                in_block_comment = False
                i += 2
            else:
                i += 1
            continue
        if in_squote:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == "'":
                in_squote = False
            i += 1
            continue
        if in_dquote:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_dquote = False
            i += 1
            continue

        if ch == "/" and nxt == "/":
            in_line_comment = True
            i += 2
            continue
        if ch == "/" and nxt == "*":
            in_block_comment = True
            i += 2
            continue
        if ch == "'":
            in_squote = True
            i += 1
            continue
        if ch == '"':
            in_dquote = True
            i += 1
            continue

        if ch == "(":
            depth += 1
        elif ch == ")":
            depth -= 1
            if depth == 0:
                close_paren = i
                break
        i += 1

    if close_paren is None:
        raise ValueError("Unbalanced parentheses in method signature")

    brace = s.find("{", close_paren)
    if brace == -1:
        raise ValueError("No '{' found for method body")
    body = brace_slice(s, brace)
    return (s[start:brace] + body).strip("\n\r")


def replace_paragraph_text(paragraph, new_text: str):
    # clear runs
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)
    paragraph.add_run(new_text)


def main():
    doc = Document(INP)
    start, end = find_chapter_range(doc, 4)
    if start is None:
        raise SystemExit("Chapter 4 not found")

    ch4_paras = doc.paragraphs[start:(end or len(doc.paragraphs))]
    code_paras = [p for p in ch4_paras if p.style and p.style.name == 'Source Code' and (p.text or '').strip()]

    if len(code_paras) != 38:
        raise SystemExit(f"Expected 38 code paragraphs in chapter 4, got {len(code_paras)}")

    repl = [None] * 38

    # 1-2 environment config snippets
    repl[0] = extract_pom_dependencies()
    repl[1] = extract_package_dependencies()

    # 3-5 user module
    repl[2] = extract_java_method(ACCOUNT_CTRL, 'register')
    repl[3] = extract_java_method(ACCOUNT_CTRL, 'login')
    repl[4] = extract_java_method(ACCOUNT_CTRL, 'updateUserInfo')

    # 6-8 meeting module
    repl[5] = extract_java_method(MEETING_INFO_CTRL, 'quickMeeting')
    repl[6] = extract_java_method(MEETING_INFO_CTRL, 'joinMeeting')
    repl[7] = extract_java_method(MEETING_RESERVE_CTRL, 'createMeetingReserve')

    # 9-11 webrtc
    repl[8] = extract_js_method(WEBRTC, r"async\s+createPeerConnection\s*\(")
    repl[9] = extract_js_method(WEBRTC, r"async\s+createAndSendOffer\s*\(")
    repl[10] = extract_js_method(WEBRTC, r"async\s+startScreenShare\s*\(")

    # 12-13 chat
    repl[11] = extract_java_method(CHAT_CTRL, 'sendMessage')
    repl[12] = extract_java_method(CHAT_CTRL, 'loadHistory')

    # 14-15 contacts
    repl[13] = extract_java_method(USER_CONTACT_CTRL, 'contactApply')
    repl[14] = extract_java_method(USER_CONTACT_CTRL, 'dealWithApply')

    # 16 notification creation example
    # Use createMeetingInviteNotification to show constructing UserNotification
    notif_service = ROOT / "src/main/java/com/easymeeting/service/impl/UserNotificationServiceImpl.java"
    repl[15] = extract_java_method(notif_service, 'createMeetingInviteNotification')

    # 17 notification list
    repl[16] = extract_java_method(NOTIFICATION_CTRL, 'loadNotificationList')

    # 18 mapper interfaces (strip package/imports)
    ui = extract_java_type(USER_INFO_MAPPER, 'UserInfoMapper', kind='interface')
    mi = extract_java_type(MEETING_INFO_MAPPER, 'MeetingInfoMapper', kind='interface')
    # Keep only interface blocks (drop package and imports if present)
    def only_interface_block(snippet: str) -> str:
        m = re.search(r"public\s+interface\b", snippet)
        return snippet[m.start():].strip() if m else snippet.strip()
    repl[17] = only_interface_block(ui) + "\n\n" + only_interface_block(mi)

    # 19 pagination query example
    repl[18] = extract_java_method(MEETING_INFO_SERVICE_IMPL, 'findListByPage')

    # 20 app main (mapper scan)
    repl[19] = extract_java_type(APP_MAIN, 'EasymeetingApplication', kind='class')

    # 21 redis config
    repl[20] = extract_java_type(REDIS_CONFIG, 'RedisConfig', kind='class')

    # 22 redisson config
    repl[21] = extract_java_type(REDISSON_CONFIG, 'RedissonConfig', kind='class')

    # 23 global interceptor annotation
    repl[22] = extract_java_type(GLOBAL_INTERCEPTOR, 'globalInterceptor', kind='@interface')

    # 24 aspect
    repl[23] = extract_java_type(GLOBAL_ASPECT, 'globalOperationAspect', kind='class')

    # 25 global exception handler
    repl[24] = extract_java_type(GLOBAL_EXCEPTION, 'AGlobalExceptionHandlerController', kind='class')

    # 26-27 websocket
    repl[25] = extract_java_type(NETTY_STARTER, 'NettyWebSocketStarter', kind='class')
    repl[26] = extract_java_type(NETTY_HANDLER, 'HandlerWebSocket', kind='class')

    # 28 redis component
    repl[27] = extract_java_type(REDIS_COMPONENT, 'RedisComponent', kind='class')

    # 29-30 minio
    repl[28] = extract_java_type(MINIO_CONFIG, 'MinioConfig', kind='class')
    repl[29] = extract_java_type(MINIO_SERVICE_IMPL, 'MinioServiceImpl', kind='class')

    # 31 router config
    repl[30] = read_text(ROUTER).strip()

    # 32 meeting page template
    repl[31] = extract_between_tags(read_text(MEETING_VUE), '<template>', '</template>')

    # 33 axios wrapper
    repl[32] = read_text(AXIOS_JS).strip()

    # 34 services excerpt (auth service)
    ssvc = read_text(SERVICES_JS)
    # take authService and meetingService header part
    m = re.search(r"export\s+const\s+authService\s*=\s*\{", ssvc)
    if not m:
        raise ValueError('authService not found')
    start_auth = ssvc.rfind('\n', 0, m.start()) + 1
    brace = ssvc.find('{', m.end()-1)
    body = brace_slice(ssvc, brace)
    repl[33] = (ssvc[start_auth:brace] + body).strip()

    # 35 element-plus login template
    repl[34] = extract_between_tags(read_text(LOGIN_VUE), '<template>', '</template>')

    # 36-37 electron
    repl[35] = '\n'.join(read_text(ELECTRON_MAIN).splitlines()[:120]).strip()
    # preload: include exposeInMainWorld block if present
    preload_text = read_text(ELECTRON_PRELOAD)
    if 'contextBridge.exposeInMainWorld' in preload_text:
        # extract from that line to end of object literal
        repl[36] = extract_js_object_literal(ELECTRON_PRELOAD, 'contextBridge.exposeInMainWorld')
    else:
        repl[36] = '\n'.join(preload_text.splitlines()[:120]).strip()

    # 38 desktop notification method (frontend)
    repl[37] = extract_js_method(SETTINGS_MANAGER, r"async\s+showDesktopNotification\s*\(")

    # Apply replacements
    for i, (p, new_text) in enumerate(zip(code_paras, repl), start=1):
        if not new_text or not new_text.strip():
            raise ValueError(f"Empty replacement for block {i}")
        replace_paragraph_text(p, new_text)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == '__main__':
    main()
