from pathlib import Path
from docx import Document

DOCX = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正.docx")

doc = Document(DOCX)

# chapter 4 range
start=end=None
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章')):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章')):
        end=i
        break

ch4 = doc.paragraphs[start:(end or len(doc.paragraphs))]

blocks=[]
for rel_i,p in enumerate(ch4):
    if not (p.style and p.style.name=='Source Code'):
        continue
    text=p.text or ''
    if not text.strip():
        continue
    lines=[ln.rstrip() for ln in text.splitlines()]
    first=''
    for ln in lines:
        if ln.strip():
            first=ln.strip()
            break
    blocks.append({'abs_index': start+rel_i, 'rel_index': rel_i, 'first': first, 'lines': len(lines)})

print('code blocks',len(blocks))
for i,b in enumerate(blocks, start=1):
    print(f"{i:02d} para#{b['abs_index']} lines={b['lines']} first={b['first'][:90]}")
