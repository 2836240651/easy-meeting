import re
import string
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SRC = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节.docx")
DST = SRC.with_name(SRC.stem + "_第5章表题字号字体修正.docx")

CHAPTER = 5
SONG = "宋体"
TNR = "Times New Roman"
SIZE = Pt(10.5)  # 五号

PUNCT = set("，。；：？！、,.!?;:（）()【】[]《》<>“”‘’·/\\|-—–－")

SECTION_RE = re.compile(r"^5\.(\d+)(?:\.(\d+))?\s+(.+?)\s*$")
CAPTION_NUM_RE = re.compile(r"^(续表)?\s*表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")
LABEL_RE = re.compile(r"表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def is_chapter_heading(text: str, chapter: int) -> bool:
    t = text.strip()
    return t.startswith(f"第{chapter}章") or t.startswith("第" + "一二三四五六七八九十"[chapter-1] + "章")


def set_run_fonts(run, east_asia: str, ascii_font: str, size=SIZE, bold=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    run.font.name = ascii_font


def sanitize_title(title: str) -> str:
    t = title.strip()
    # remove leading enumerations like （1） or (1)
    t = re.sub(r"^（\d+）\s*", "", t)
    t = re.sub(r"^\(\d+\)\s*", "", t)
    t = re.sub(r"^\d+[\.)、]\s*", "", t)
    # remove trailing colon-like
    t = re.sub(r"[：:]\s*$", "", t)
    # strip punctuation characters inside
    t = "".join(ch for ch in t if ch not in PUNCT)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def paragraph_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def insert_paragraph_before(tbl: Table, text: str) -> Paragraph:
    tbl_elm = tbl._element
    p = OxmlElement('w:p')
    tbl_elm.addprevious(p)
    para = Paragraph(p, tbl._parent)
    para.add_run(text)
    return para


def make_caption(paragraph: Paragraph, label: str, title: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # clear existing runs
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)

    # construct runs: 表 + label + space + title
    r1 = paragraph.add_run("表")
    set_run_fonts(r1, SONG, TNR, bold=True)

    r2 = paragraph.add_run(label)
    set_run_fonts(r2, TNR, TNR, bold=True)

    r3 = paragraph.add_run(" " + title)
    set_run_fonts(r3, SONG, TNR, bold=True)

    # keep caption with table
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = pPr.find(qn('w:keepNext'))
    if keepNext is None:
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)


def format_table_fonts(tbl: Table):
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = trPr.find(qn('w:cantSplit'))
        if cant is None:
            cant = OxmlElement('w:cantSplit')
            trPr.append(cant)

    # mark header row repeat
    if tbl.rows:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        hdr = trPr.find(qn('w:tblHeader'))
        if hdr is None:
            hdr = OxmlElement('w:tblHeader')
            trPr.append(hdr)

    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    # keep bold/italic as-is, just normalize fonts and size
                    set_run_fonts(run, SONG, TNR, size=SIZE, bold=None)


def update_in_text_refs(paragraphs, mapping):
    if not mapping:
        return 0
    count = 0
    for p in paragraphs:
        text = p.text
        if not text:
            continue
        new = text
        for old, new_label in mapping.items():
            # handle forms like 表5-1 / 表 5-1 / 表-5.3.1
            new = re.sub(rf"表\s*[-–—－]?\s*{re.escape(old)}", f"表{new_label}", new)
        if new != text:
            # replace paragraph text while preserving run formatting as much as possible:
            # simplest: set full text in one run and then normalize fonts.
            for r in list(p.runs):
                p._element.remove(r._element)
            r = p.add_run(new)
            set_run_fonts(r, SONG, TNR, size=SIZE, bold=None)
            count += 1
    return count


doc = Document(SRC)
blocks = list(iter_block_items(doc))

in_ch5 = False
current_section_title = ""
ch5_tables = []  # list of (block_idx, Table)
ch5_paragraphs = []

for i, b in enumerate(blocks):
    if isinstance(b, Paragraph):
        t = paragraph_text(b)
        if is_chapter_heading(t, CHAPTER):
            in_ch5 = True
        elif is_chapter_heading(t, CHAPTER + 1) and in_ch5:
            in_ch5 = False
        if in_ch5:
            ch5_paragraphs.append(b)
            m = SECTION_RE.match(t)
            if m:
                current_section_title = m.group(3).strip()
    elif isinstance(b, Table) and in_ch5:
        ch5_tables.append((i, b))

print(f"Found chapter {CHAPTER} tables: {len(ch5_tables)}")

# Build old label mapping from existing captions (immediate previous paragraph)
old_to_new = {}

def extract_old_label(caption_text: str):
    ct = caption_text.strip()
    if not ct:
        return None
    m = CAPTION_NUM_RE.match(ct)
    if m:
        return m.group(2)
    # also allow something like 表5-1 ...
    m2 = LABEL_RE.search(ct)
    if m2:
        return m2.group(1)
    return None

# Determine titles and captions
changes = []
for seq, (tbl_block_idx, tbl) in enumerate(ch5_tables, start=1):
    new_label = f"{CHAPTER}-{seq}"

    # locate immediate previous non-empty paragraph block
    prev_para = None
    prev_text = ""
    j = tbl_block_idx - 1
    while j >= 0:
        if isinstance(blocks[j], Paragraph):
            t = paragraph_text(blocks[j])
            if t:
                prev_para = blocks[j]
                prev_text = t
                break
        elif isinstance(blocks[j], Table):
            break
        j -= 1

    caption_para = None
    caption_title = ""
    if prev_para is not None and (prev_text.startswith('表') or prev_text.startswith('续表')):
        caption_para = prev_para
        old = extract_old_label(prev_text)
        if old:
            old_to_new[old] = new_label
        # get title remainder
        m = CAPTION_NUM_RE.match(prev_text)
        if m:
            caption_title = m.group(3)
        else:
            # remove leading 表...digits
            caption_title = re.sub(r"^(续表)?\s*表\s*[-–—－]?\s*[0-9]+(?:[\.-][0-9]+)*", "", prev_text)
        caption_title = sanitize_title(caption_title)

    if not caption_title:
        # derive title from previous text or section heading
        candidate = ""
        if prev_text and not (prev_text.startswith('表') or prev_text.startswith('续表')):
            candidate = sanitize_title(prev_text)
        if candidate in ("测试用例", "测试用例表", "性能测试结果", "安全测试用例", "压力测试结果", "主要接口测试结果", "测试场景", "测试场景设计", "测试结果"):
            candidate = ""

        # get first row header concat
        header = ""
        try:
            if tbl.rows:
                header = " ".join(c.text.strip().replace('\n',' ') for c in tbl.rows[0].cells)
        except Exception:
            header = ""

        if not candidate and current_section_title:
            candidate = sanitize_title(current_section_title)

        if '用例编号' in header or ('测试步骤' in header and '预期结果' in header):
            caption_title = (candidate or "测试") + "测试用例表"
        elif '响应时间' in header or 'TPS' in header:
            caption_title = (candidate or "接口性能测试") + "结果表"
        elif '浏览器' in header or '桌面应用' in header or 'Web端' in header:
            caption_title = (candidate or "兼容性测试") + "结果表"
        elif '并发用户数' in header and '平均响应时间' in header:
            caption_title = (candidate or "压力测试") + "结果表"
        elif '配置说明' in header or '设备类型' in header:
            caption_title = (candidate or "测试环境") + "配置表"
        elif '接口路径' in header and '方法' in header:
            # if this was a continuation table, differentiate
            if prev_text.startswith('续表'):
                caption_title = (candidate or "接口测试") + "续表"
            else:
                caption_title = (candidate or "接口测试") + "结果表"
        else:
            caption_title = (candidate or "表")
            if not caption_title.endswith('表'):
                caption_title += "表"

        caption_title = sanitize_title(caption_title)

    # insert caption paragraph if missing
    if caption_para is None:
        caption_para = insert_paragraph_before(tbl, "")
        # update blocks list? not needed for following tables because insertion occurs before current table.
        changes.append((seq, 'insert'))
    else:
        changes.append((seq, 'update'))

    # If it was a continuation caption starting with 续表, mark title as 续表 if not already
    if prev_text.startswith('续表') and '续表' not in caption_title:
        caption_title = sanitize_title(caption_title + '续表')

    make_caption(caption_para, new_label, caption_title)

    # format table fonts
    format_table_fonts(tbl)

# Update in-text references within chapter 5
ref_updates = update_in_text_refs(ch5_paragraphs, old_to_new)

print(f"Caption ops: {len(changes)}; in-text ref paragraphs updated: {ref_updates}")
print(f"Old->New labels updated in text: {old_to_new}")

doc.save(DST)
print(f"Saved: {DST}")
