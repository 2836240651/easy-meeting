from pathlib import Path
from docx import Document

DOCX = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正.docx")
doc = Document(DOCX)

# ch4 range
start=end=None
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章')):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章')):
        end=i
        break

ch4 = doc.paragraphs[start:(end or len(doc.paragraphs))]

# collect source code paras
indices=[]
for rel_i,p in enumerate(ch4):
    if p.style and p.style.name=='Source Code' and (p.text or '').strip():
        indices.append(rel_i)

print('source code blocks',len(indices))

for n, rel_i in enumerate(indices, start=1):
    abs_i = start + rel_i
    p = ch4[rel_i]
    lines=[ln.rstrip() for ln in (p.text or '').splitlines()]
    first3=[ln for ln in lines if ln.strip()][:3]
    # find nearest previous heading (Heading 2/3) within 10 paras
    heading=None
    for back in range(1,25):
        j=rel_i-back
        if j<0: break
        pp=ch4[j]
        if pp.style and pp.style.name in ('Heading 2','Heading 3'):
            ht=(pp.text or '').strip()
            if ht:
                heading=ht
                break
    print(f"\n{n:02d} abs#{abs_i} heading={heading}")
    for k,ln in enumerate(first3, start=1):
        print(f"  L{k}: {ln[:140]}")