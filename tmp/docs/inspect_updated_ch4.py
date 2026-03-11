from pathlib import Path
from docx import Document

DOCX = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正_第4章代码按项目更新.docx")

doc = Document(DOCX)

# find chapter 4
start=end=None
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章')):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章')):
        end=i
        break

ch4 = doc.paragraphs[start:(end or len(doc.paragraphs))]
code=[p for p in ch4 if p.style and p.style.name=='Source Code' and (p.text or '').strip()]
print('code blocks',len(code))

for idx,p in enumerate(code, start=1):
    lines=[ln for ln in (p.text or '').splitlines() if ln.strip()]
    print(f"\n{idx:02d} first: {lines[0][:100]}")
    if len(lines)>1:
        print(f"   second: {lines[1][:100]}")