import re
from pathlib import Path
from docx import Document

DOCX = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正.docx")

doc = Document(DOCX)

# find chapter 4 start/end indices
start=None
end=None
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章') or t=='第4章' or t=='第四章'):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章') or t=='第5章' or t=='第五章'):
        end=i
        break

print('ch4 start',start,'end',end)

if start is not None:
    ch4_paras = doc.paragraphs[start:(end if end is not None else len(doc.paragraphs))]
    # print headings-like lines and any lines containing '代码' or '如图' '如下'
    head_pat=re.compile(r'^4\.(\d+)(?:\.(\d+))?\s+')
    hits=[]
    for p in ch4_paras:
        t=(p.text or '').strip()
        if not t:
            continue
        if t.startswith('4.') or t.startswith('（') or '代码' in t or '如图' in t or '如下' in t or '见' in t:
            if head_pat.match(t) or t.startswith('第4章') or '代码' in t:
                hits.append(t)
    print('\n'.join(hits[:200]))

    # also count existing code-style paragraphs (using style name contains 'Code' or monospaced?)
    styles=set()
    for p in ch4_paras:
        if p.style is not None:
            styles.add(p.style.name)
    print('styles in ch4 sample:', sorted(list(styles))[:30])
