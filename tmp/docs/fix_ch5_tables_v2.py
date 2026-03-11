import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SRC = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节.docx")
DST = SRC.with_name(SRC.stem + "_第5章表题字号字体修正_v2.docx")

CHAPTER = 5
SONG = "宋体"
TNR = "Times New Roman"
SIZE = Pt(10.5)  # 五号

PUNCT = set("，。；：？！、,.!?;:（）()【】[]《》<>“”‘’·/\\|—–－")

SECTION_RE = re.compile(r"^5\.(\d+)(?:\.(\d+))?\s+(.+?)\s*$")
CAPTION_RE = re.compile(r"^(续表)?\s*表\s*[-–—－]?\s*([0-9]+(?:[\.-][0-9]+)*)\s*(.*)$")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def is_chapter_heading(text: str, chapter: int) -> bool:
    t = text.strip()
    cn = "一二三四五六七八九十"[chapter-1]
    return t.startswith(f"第{chapter}章") or t.startswith(f"第{cn}章")


def set_run_fonts(run, east_asia: str, ascii_font: str, size=SIZE, bold=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    run.font.name = ascii_font


def sanitize_title(title: str) -> str:
    t = (title or "").strip()
    # remove leading section numbering like 5.4.1
    t = re.sub(r"^[0-9]+(?:\.[0-9]+)+\s*", "", t)
    t = re.sub(r"^5\.[0-9]+(?:\.[0-9]+)?\s*", "", t)
    # remove leading enumerations like （1） or (1)
    t = re.sub(r"^（\d+）\s*", "", t)
    t = re.sub(r"^\(\d+\)\s*", "", t)
    t = re.sub(r"^\d+[\.)、]\s*", "", t)
    # remove trailing colon-like
    t = re.sub(r"[：:]\s*$", "", t)
    # strip punctuation characters inside
    t = "".join(ch for ch in t if ch not in PUNCT)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def insert_paragraph_before(tbl: Table) -> Paragraph:
    tbl_elm = tbl._element
    p = OxmlElement('w:p')
    tbl_elm.addprevious(p)
    return Paragraph(p, tbl._parent)


def make_caption(paragraph: Paragraph, label: str, title: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # clear existing runs
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)

    r1 = paragraph.add_run("表")
    set_run_fonts(r1, SONG, TNR, bold=True)

    r2 = paragraph.add_run(label)
    set_run_fonts(r2, TNR, TNR, bold=True)

    r3 = paragraph.add_run(" " + title)
    set_run_fonts(r3, SONG, TNR, bold=True)

    # keep caption with next (table)
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = pPr.find(qn('w:keepNext'))
    if keepNext is None:
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)


def format_table_fonts(tbl: Table):
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = trPr.find(qn('w:cantSplit'))
        if cant is None:
            cant = OxmlElement('w:cantSplit')
            trPr.append(cant)

    if tbl.rows:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        hdr = trPr.find(qn('w:tblHeader'))
        if hdr is None:
            hdr = OxmlElement('w:tblHeader')
            trPr.append(hdr)

    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_fonts(run, SONG, TNR, size=SIZE, bold=None)


def extract_caption_parts(text: str):
    m = CAPTION_RE.match((text or "").strip())
    if not m:
        return None, None, None
    is_cont = bool(m.group(1))
    old_label = m.group(2)
    rest = m.group(3)
    return is_cont, old_label, rest


def choose_title(record):
    prev_text = record['prev_text']
    section_title = record['section_title']
    header = record['header']

    # prefer existing caption title
    if record['had_caption']:
        return record['caption_title']

    prev_candidate = sanitize_title(prev_text)
    if prev_candidate and (len(prev_candidate) <= 14) and not any(w in prev_candidate for w in ["以下", "进行", "如下", "本节", "主要", "包括"]):
        # common labels in this doc that are too generic
        if prev_candidate not in {"测试用例", "测试场景", "测试结果", "性能测试结果", "安全测试用例", "主要接口测试结果", "压力测试结果"}:
            # normalize some known cases
            if "环境" in prev_candidate and not prev_candidate.endswith("配置"):
                return prev_candidate + "配置表"
            if prev_candidate.endswith("结果"):
                return prev_candidate + "表"
            if not prev_candidate.endswith("表"):
                return prev_candidate + "表"
            return prev_candidate

    base = sanitize_title(section_title) or "表"

    if '用例编号' in header or ('测试步骤' in header and '预期结果' in header):
        return sanitize_title(base + "测试用例表")
    if '接口路径' in header and '方法' in header:
        if record.get('is_continuation'):
            return sanitize_title("RESTful API接口测试续表")
        return sanitize_title("RESTful API接口测试")
    if '响应时间' in header or 'TPS' in header:
        return sanitize_title(base + "结果表")
    if '桌面应用' in header or 'Web端' in header or '浏览器' in header:
        return sanitize_title(base + "结果表")
    if '并发用户数' in header and '错误率' in header:
        return sanitize_title(base + "结果表")
    if '配置说明' in header or '设备类型' in header:
        if '客户端' in prev_text:
            return sanitize_title("客户端测试环境配置表")
        return sanitize_title("测试环境配置表")

    if not base.endswith('表'):
        base += '表'
    return sanitize_title(base)


def update_single_paragraph_text(p: Paragraph, new_text: str):
    # Replace the paragraph content with a single run; normalize fonts.
    for r in list(p.runs):
        p._element.remove(r._element)
    r = p.add_run(new_text)
    set_run_fonts(r, SONG, TNR, size=SIZE, bold=None)


doc = Document(SRC)
blocks = list(iter_block_items(doc))

in_ch5 = False
section_title = ""
records = []
ch5_paragraphs = []

for idx, b in enumerate(blocks):
    if isinstance(b, Paragraph):
        t = (b.text or "").strip()
        if is_chapter_heading(t, CHAPTER):
            in_ch5 = True
        elif is_chapter_heading(t, CHAPTER + 1) and in_ch5:
            in_ch5 = False
        if in_ch5:
            ch5_paragraphs.append(b)
            m = SECTION_RE.match(t)
            if m:
                section_title = m.group(3).strip()
    elif isinstance(b, Table) and in_ch5:
        # immediate prev non-empty paragraph (stop at previous table)
        prev_para = None
        prev_text = ""
        j = idx - 1
        while j >= 0:
            if isinstance(blocks[j], Paragraph):
                tx = (blocks[j].text or "").strip()
                if tx:
                    prev_para = blocks[j]
                    prev_text = tx
                    break
            elif isinstance(blocks[j], Table):
                break
            j -= 1

        header = ""
        if b.rows:
            header = " ".join(c.text.strip().replace('\n',' ') for c in b.rows[0].cells)

        had_caption = bool(prev_text.startswith('表') or prev_text.startswith('续表'))
        is_continuation = prev_text.startswith('续表')

        old_label = None
        caption_title = ""
        if had_caption:
            _, old_label, rest = extract_caption_parts(prev_text)
            caption_title = sanitize_title(rest)
            if is_continuation and '续表' not in caption_title:
                caption_title = sanitize_title(caption_title + "续表")

        records.append({
            'table': b,
            'block_idx': idx,
            'prev_para': prev_para,
            'prev_text': prev_text,
            'section_title': section_title,
            'header': header,
            'had_caption': had_caption,
            'is_continuation': is_continuation,
            'old_label': old_label,
            'caption_title': caption_title,
        })

print(f"Chapter {CHAPTER} tables: {len(records)}")

old_to_new = {}
caption_paras = set()

for seq, rec in enumerate(records, start=1):
    new_label = f"{CHAPTER}-{seq}"

    if rec['old_label']:
        old_to_new[rec['old_label']] = new_label

    title = choose_title({
        **rec,
        'caption_title': rec['caption_title'],
        'is_continuation': rec['is_continuation'],
    })

    # ensure no punctuation and no trailing punctuation already handled; ensure at least something
    title = sanitize_title(title) or "表"

    if rec['had_caption'] and rec['prev_para'] is not None:
        caption_para = rec['prev_para']
    else:
        caption_para = insert_paragraph_before(rec['table'])

    make_caption(caption_para, new_label, title)
    caption_paras.add(caption_para._element)

    format_table_fonts(rec['table'])

# Update in-text references (excluding caption paragraphs) with boundary-safe replacement
if old_to_new:
    pattern_items = list(old_to_new.items())
    for p in ch5_paragraphs:
        if p._element in caption_paras:
            continue
        text = p.text or ""
        new = text
        for old, new_label in pattern_items:
            # ensure old label not followed by a digit (avoid 5-1 matching 5-11)
            new = re.sub(rf"表\s*[-–—－]?\s*{re.escape(old)}(?!\d)", f"表{new_label}", new)
        if new != text:
            update_single_paragraph_text(p, new)


doc.save(DST)
print(f"Saved: {DST}")
print(f"Old->New mapping: {old_to_new}")
