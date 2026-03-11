from docx import Document
from pathlib import Path

p=Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正_第4章代码按项目更新.docx")
d=Document(p)

# chapter markers sample
chs=[]
for para in d.paragraphs:
    t=(para.text or '').strip()
    if t.startswith('第') and '章' in t[:6] and len(t)<30:
        chs.append(t)
print('chapter markers sample:', chs[:12])

# chapter 4 range
start=end=None
for i,para in enumerate(d.paragraphs):
    t=(para.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章')):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章')):
        end=i
        break
ch4=d.paragraphs[start:(end or len(d.paragraphs))]
code=[p for p in ch4 if p.style and p.style.name=='Source Code' and (p.text or '').strip()]
print('ch4 code blocks:', len(code))

# marker paragraphs followed by code
checks=[('用户注册功能','register'),('WebRTC 连接建立','createPeerConnection')]
for mk, expect in checks:
    ok=False
    for i,p2 in enumerate(ch4):
        tx=(p2.text or '')
        if mk in tx and '代码如下' in tx:
            j=i+1
            while j<len(ch4) and not (ch4[j].text or '').strip():
                j+=1
            nxt=ch4[j]
            first=(nxt.text or '').splitlines()[0] if (nxt.text or '').splitlines() else ''
            ok = (nxt.style and nxt.style.name=='Source Code' and expect in (nxt.text or ''))
            print(mk,'->', 'OK' if ok else 'NOT_OK', 'next_style', nxt.style.name if nxt.style else None, 'first', first[:80])
            break
    if not ok:
        pass