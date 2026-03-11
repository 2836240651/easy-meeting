import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

INP = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正_第4章代码按项目更新.docx")
OUT = INP.with_name(INP.stem + "_注释小四宋体.docx")

SONG = "宋体"
SMALL4 = Pt(12)  # 小四


def find_chapter_range(doc, chapter: int):
    start = end = None
    cn = "一二三四五六七八九十"
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if start is None and (t.startswith(f"第{chapter}章") or t.startswith(f"第{cn[chapter-1]}章")):
            start = i
        elif start is not None and end is None and (t.startswith(f"第{chapter+1}章") or t.startswith(f"第{cn[chapter]}章")):
            end = i
            break
    return start, end


def set_run_song_small4(run):
    run.font.name = SONG
    run.font.size = SMALL4
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), SONG)
    rFonts.set(qn('w:ascii'), SONG)
    rFonts.set(qn('w:hAnsi'), SONG)


def find_line_comment_pos(line: str):
    # Find '//' outside of strings, ignoring '://'
    in_s = False
    in_d = False
    esc = False
    i = 0
    while i < len(line) - 1:
        ch = line[i]
        nxt = line[i + 1]

        if esc:
            esc = False
            i += 1
            continue

        if in_s:
            if ch == '\\':
                esc = True
            elif ch == "'":
                in_s = False
            i += 1
            continue

        if in_d:
            if ch == '\\':
                esc = True
            elif ch == '"':
                in_d = False
            i += 1
            continue

        if ch == "'":
            in_s = True
            i += 1
            continue
        if ch == '"':
            in_d = True
            i += 1
            continue

        if ch == '/' and nxt == '/':
            # ignore URLs like http:// or https://
            if i > 0 and line[i - 1] == ':':
                i += 2
                continue
            return i

        i += 1

    return None


def rewrite_paragraph_with_comment_font(p):
    text = p.text or ""
    if '//' not in text:
        return False

    lines = text.splitlines()
    new_runs = []

    for li, line in enumerate(lines):
        cpos = find_line_comment_pos(line)
        if cpos is None:
            new_runs.append(('code', line))
        else:
            # keep '//' in code font; comment after it becomes Song small4
            new_runs.append(('code', line[:cpos + 2]))
            new_runs.append(('comment', line[cpos + 2:]))

        if li != len(lines) - 1:
            new_runs.append(('code', "\n"))

    # clear existing runs
    for r in list(p.runs):
        p._element.remove(r._element)

    for kind, chunk in new_runs:
        if chunk == "":
            continue
        r = p.add_run(chunk)
        if kind == 'comment':
            set_run_song_small4(r)

    return True


def main():
    doc = Document(INP)
    start, end = find_chapter_range(doc, 4)
    if start is None:
        raise SystemExit('Chapter 4 not found')

    changed = 0
    scanned = 0
    for p in doc.paragraphs[start:(end or len(doc.paragraphs))]:
        if not (p.style and p.style.name == 'Source Code'):
            continue
        scanned += 1
        if rewrite_paragraph_with_comment_font(p):
            changed += 1

    doc.save(OUT)
    print(f"Scanned source-code paragraphs: {scanned}; Updated: {changed}")
    print(f"Saved: {OUT}")


if __name__ == '__main__':
    main()