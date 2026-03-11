import re
from pathlib import Path
from docx import Document

DOCX = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正_第6章更新_精简技术细节_第5章表题字号字体修正_最终_第3章同步修正.docx")

doc = Document(DOCX)

# chapter 4 range
start=end=None
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start is None and (t.startswith('第4章') or t.startswith('第四章')):
        start=i
    elif start is not None and end is None and (t.startswith('第5章') or t.startswith('第五章')):
        end=i
        break

ch4 = doc.paragraphs[start:(end or len(doc.paragraphs))]

markers=[]
for idx,p in enumerate(ch4):
    t=(p.text or '').strip()
    if not t:
        continue
    if '代码如下' in t or '核心实现代码' in t or '核心代码' in t or t.endswith('如下：') and '代码' in t:
        markers.append((idx, t))

print('markers',len(markers))
for mi,(idx,t) in enumerate(markers, start=1):
    print('\n#',mi,'para',start+idx,':',t)
    for j in range(1,6):
        if idx+j >= len(ch4):
            break
        nt=(ch4[idx+j].text or '').rstrip()
        if nt.strip():
            print('  next',j,':',nt[:160])

# also locate Source Code style paragraphs inside ch4
src_paras=[]
for idx,p in enumerate(ch4):
    if (p.style and p.style.name=='Source Code') and (p.text or '').strip():
        src_paras.append((start+idx, (p.text or '')[:120]))
print('\nsource-code paragraphs with text:',len(src_paras))
for i,(pi,tx) in enumerate(src_paras[:30], start=1):
    print(i, pi, tx.replace('\n','\\n'))
