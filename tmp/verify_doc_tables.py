from docx import Document


path = r"C:\Users\13377\Desktop\论文相关\1.0.1_v1.0.2.docx"
doc = Document(path)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("表") or text.startswith("续表"):
        print("TITLE", i, text)
        for run in p.runs[:6]:
            size = run.font.size.pt if run.font.size else None
            print(" ", repr(run.text), run.font.name, size, run.bold)
        break

cell = doc.tables[1].cell(0, 0)
paragraph = cell.paragraphs[0]
print("CELL", paragraph.text)
for run in paragraph.runs[:6]:
    size = run.font.size.pt if run.font.size else None
    print(" ", repr(run.text), run.font.name, size, run.bold)
