import re
from pathlib import Path
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新.docx")
if not in_path.exists():
    in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")

out_path = in_path.with_name(in_path.stem + '_5_2_1表格三线.docx')

doc = docx.Document(str(in_path))
body = doc._body._body

# find section bounds
start_idx = None
end_idx = None
for p in doc.paragraphs:
    if p.text.strip() == '5.2.1 用户模块测试':
        start_idx = body.index(p._element)
        break
if start_idx is None:
    raise SystemExit('5.2.1 not found')
for p in doc.paragraphs:
    if body.index(p._element) > start_idx and p.text.strip().startswith('5.2.2'):
        end_idx = body.index(p._element)
        break
if end_idx is None:
    raise SystemExit('5.2.2 not found')

# collect existing table numbers in chapter 5
max_table_num = 0
for p in doc.paragraphs:
    m = re.search(r'表5-(\d+)', p.text)
    if m:
        max_table_num = max(max_table_num, int(m.group(1)))

# find tables between start and end in body
section_tbl_idxs = []
for idx in range(start_idx + 1, end_idx):
    elem = body[idx]
    if elem.tag.endswith('}tbl'):
        section_tbl_idxs.append(idx)

if not section_tbl_idxs:
    print('no tables found in 5.2.1')

# build list of new table numbers
new_nums = []
for _ in section_tbl_idxs:
    max_table_num += 1
    new_nums.append(max_table_num)

# update the （1）测试用例 paragraph to include references
ref_text = ''
if new_nums:
    refs = '、'.join([f'表5-{n}' for n in new_nums])
    ref_text = f'（1）测试用例。如{refs}所示。'

for p in doc.paragraphs:
    if body.index(p._element) > start_idx and body.index(p._element) < end_idx:
        if p.text.strip() == '（1）测试用例':
            if ref_text:
                p.text = ref_text
            break

# helper to set caption font

def set_run_font(run, font_name, size=10.5, bold=True, east_asia=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    if east_asia:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)

# apply three-line style

def apply_three_line_table(table):
    tbl = table._element
    tblPr = tbl.tblPr
    # remove borders by setting nil
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        tblBorders.append(elem)
    tblPr.append(tblBorders)

    def set_row_border(row, top=False, bottom=False):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trBorders = OxmlElement('w:trBorders')
        for edge, enable in (('top', top), ('bottom', bottom)):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single' if enable else 'nil')
            elem.set(qn('w:sz'), '8')
            elem.set(qn('w:color'), '000000')
            trBorders.append(elem)
        trPr.append(trBorders)

    if table.rows:
        set_row_border(table.rows[0], top=True, bottom=True)
        set_row_border(table.rows[-1], bottom=True)

# insert captions and apply table style
for idx, tbl_idx in enumerate(section_tbl_idxs):
    table_num = new_nums[idx]
    # find table object by index among all tables
    # map body index to table object
    table = None
    # locate by matching element
    for t in doc.tables:
        if t._tbl is body[tbl_idx]:
            table = t
            break
    if table is None:
        continue

    # caption text: 用户模块测试用例表 (no punctuation)
    caption_text = f'表5-{table_num} 用户模块测试用例表'

    cap = doc.add_paragraph('')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = cap.add_run('表')
    run2 = cap.add_run(f'5-{table_num}')
    run3 = cap.add_run(' 用户模块测试用例表')
    set_run_font(run1, '宋体', east_asia='宋体')
    set_run_font(run2, 'Times New Roman', east_asia='Times New Roman')
    set_run_font(run3, '宋体', east_asia='宋体')

    # insert caption before table element
    body.remove(cap._element)
    body.insert(tbl_idx, cap._element)

    # apply three-line table style
    apply_three_line_table(table)

# save
_doc = doc
_doc.save(str(out_path))
print('updated', out_path)
