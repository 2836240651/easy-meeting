from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全.docx")
OUT = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v5.docx")

ANCHOR_KEYS = (
    "\u4e3b\u8981\u7814\u7a76\u5185\u5bb9\u5305\u62ec",
    "\u7814\u7a76\u76ee\u6807\u662f",
    "\u4ee5\u4e0b\u6838\u5fc3\u7279\u6027",
    "\u4e3b\u8981\u5e94\u7528\u4e8e\u4ee5\u4e0b\u51e0\u4e2a\u65b9\u9762",
    "\u6838\u5fc3API\u5305\u62ec",
    "\u901a\u4fe1\u8fc7\u7a0b\u4e3b\u8981\u5305\u62ec\u4ee5\u4e0b\u51e0\u4e2a\u9636\u6bb5",
    "\u5177\u4f53\u5e94\u7528\u5305\u62ec",
    "\u4e3b\u8981\u7279\u70b9\u5305\u62ec",
    "\u4e3b\u8981\u5e94\u7528\u573a\u666f\u5305\u62ec",
    "\u6838\u5fc3\u7279\u6027\u5305\u62ec",
    "\u7cfb\u7edf\u91c7\u7528\u4ee5\u4e0b\u6d4b\u8bd5\u65b9\u6cd5",
    "\u4ee5\u4e0b\u73af\u5883\u4e2d\u8fdb\u884c",
    "\u53ef\u4ee5\u4ece\u4ee5\u4e0b\u51e0\u4e2a\u65b9\u9762\u8fdb\u884c\u6539\u8fdb\u548c\u6269\u5c55",
    "\u4e3b\u8981\u5de5\u4f5c\u548c\u521b\u65b0\u70b9\u5305\u62ec",
    "\u4e0d\u8db3\u4e4b\u5904",
)

HEADING_STYLES = {
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "\u7ae0\u6807\u9898",
    "\u8282\u6807\u9898",
    "\u6761\u6807\u9898",
}

COLON_CHARS = ("\uff1a", ":")


def is_heading(paragraph):
    return paragraph.style.name in HEADING_STYLES


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")


def style_item_paragraph(paragraph):
    paragraph.style = "Body Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run)


def normalize_text(text):
    return re.sub(r"^\(\d+\)\s*", "", text.strip())


def has_colon(text):
    return any(ch in text for ch in COLON_CHARS)


def is_candidate_title(text):
    text = normalize_text(text)
    if not text:
        return False
    if len(text) > 30:
        return False
    if text.endswith("\u3002"):
        return False
    if text.startswith("\u8868"):
        return False
    return True


def join_title_body(title, body):
    title = normalize_text(title)
    body = re.sub(r"^\(\d+\)\s*", "", body.strip())
    body = body.lstrip("\uff1a:")
    return f"{title}\uff1a{body}"


doc = Document(str(SRC))

i = 0
while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text and any(key in text for key in ANCHOR_KEYS):
        j = i + 1
        n = 1
        while j < len(doc.paragraphs):
            cur = doc.paragraphs[j]
            cur_text = cur.text.strip()
            if not cur_text:
                j += 1
                continue
            if is_heading(cur):
                break
            if cur_text.startswith("\u8868"):
                break

            next_para = doc.paragraphs[j + 1] if j + 1 < len(doc.paragraphs) else None
            next_text = next_para.text.strip() if next_para is not None else ""

            if (
                is_candidate_title(cur_text)
                and next_para is not None
                and next_text
                and not is_heading(next_para)
                and not has_colon(cur_text)
            ):
                cur.text = f"({n})\t{join_title_body(cur_text, next_text)}"
                style_item_paragraph(cur)
                delete_paragraph(next_para)
                n += 1
                j += 1
                continue

            if has_colon(cur_text):
                cur.text = f"({n})\t{normalize_text(cur_text)}"
                style_item_paragraph(cur)
                n += 1
                j += 1
                continue

            break

        i = j
    else:
        i += 1

doc.save(str(OUT))
print(OUT)
