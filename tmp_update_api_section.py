from pathlib import Path
import docx

doc_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")
out_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新.docx")

doc = docx.Document(str(doc_path))
body = doc._body._body

start_idx = None
end_idx = None

for p in doc.paragraphs:
    if p.text.strip() == '3.4.2 主要接口列表':
        start_idx = body.index(p._element)
        break

if start_idx is None:
    raise SystemExit('3.4.2 not found')

for p in doc.paragraphs:
    if body.index(p._element) > start_idx and p.text.strip() == '本章小结':
        end_idx = body.index(p._element)
        break

if end_idx is None:
    raise SystemExit('end not found')

for i in range(end_idx - 1, start_idx, -1):
    body.remove(body[i])

insert_idx = start_idx + 1

lines = [
    '用户模块主要接口：',
    '/account/checkCode 或 /account/captcha 获取验证码',
    '/account/login 用户登录',
    '/account/register 用户注册',
    '/account/updateUserInfo 更新用户信息',
    '/account/updatePassword 修改密码',
    '/userInfo/getUserInfo 查询用户信息',
    '/userInfo/updateUserInfo 修改用户信息',
    '/api/settings/get 获取用户设置',
    '/api/settings/save 保存用户设置',
    '/api/settings/changePassword 修改密码（设置模块）',
    '',
    '会议模块主要接口：',
    '/meetingInfo/loadMeeting 加载会议列表',
    '/meetingInfo/quickMeeting 快速创建会议',
    '/meetingInfo/preJoinMeeting 预加入会议',
    '/meetingInfo/joinMeeting 加入会议',
    '/meetingInfo/exitMeeting 退出会议',
    '/meetingInfo/getCurrentMeeting 获取当前会议',
    '/meetingInfo/finishMeeting 结束会议',
    '/meetingInfo/loadMeetingMembers 获取会议成员',
    '/meetingInfo/inviteContact 邀请联系人',
    '/meetingMember/loadDataList 会议成员列表（管理）',
    '',
    '预约模块主要接口：',
    '/meetingReserve/createMeetingReserve 创建预约会议',
    '/meetingReserve/loadMeetingReserveList 查询预约列表',
    '/meetingReserve/getMeetingReserveDetail 查询预约详情',
    '/meetingReserve/updateMeetingReserve 修改预约会议',
    '/meetingReserve/cancelMeetingReserve 取消预约会议',
    '/meetingReserve/leaveMeetingReserve 退出预约会议',
    '/meetingReserve/getUpcomingMeetings 获取即将开始会议',
    '/meetingReserveMember/loadDataList 预约成员列表（管理）',
    '',
    '联系人模块主要接口：',
    '/userContact/searchContact 搜索联系人',
    '/userContact/contactApply 发起好友申请',
    '/userContact/dealWithApply 处理好友申请',
    '/userContact/loadContactUser 获取联系人列表',
    '/userContact/loadContactApply 获取收到的申请',
    '/userContact/loadMyApply 获取我发出的申请',
    '/userContact/delContact 删除或拉黑联系人',
    '/userContact/loadBlackList 获取黑名单',
    '/userContact/unblackContact 取消拉黑',
    '',
    '通知模块主要接口：',
    '/notification/loadNotificationList 获取通知列表',
    '/notification/getUnreadCount 获取未读数量',
    '/notification/markAsRead 标记已读',
    '/notification/markAllAsRead 全部已读',
    '/notification/loadNotificationsByCategory 分类查询',
    '/notification/loadPendingActions 待处理通知',
    '/notification/handleMeetingInvite 处理会议邀请',
    '/notification/updateActionStatus 更新操作状态',
    '',
    'AI模块主要接口：',
    '/ai/chat AI对话',
    '/ai/summary 会议摘要',
    '/ai/suggest 智能建议',
    '/ai/smartSummary 智能摘要',
    '',
    '文件与头像接口：',
    '/upload/avatar 头像上传',
    '/upload/avatarByUrl 通过URL设置头像',
    '/files/avatar/{filename} 头像访问',
    '',
    '版本更新接口：',
    '/admin/loadUpdateList 获取版本更新列表',
    '/admin/saveUpdate 新增或更新版本信息',
    '/admin/delUpdate 删除版本更新',
    '/admin/postUpdate 发布版本更新',
]

for text in lines:
    p = doc.add_paragraph(text)
    body.remove(p._element)
    body.insert(insert_idx, p._element)
    insert_idx += 1

_doc = doc
_doc.save(str(out_path))
print('updated', out_path)
