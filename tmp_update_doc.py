import re
from pathlib import Path
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sql_path = Path(r"D:\JavaPartical\easymeeting-java\easymeeting_tables_excluding_chat.sql")
doc_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")

sql = sql_path.read_text(encoding="utf-8", errors="ignore")

create_re = re.compile(r"CREATE TABLE `(?P<name>[^`]+)` \((?P<body>.*?)\) ENGINE=", re.S)
col_re = re.compile(r"^\s*`(?P<col>[^`]+)`\s+(?P<type>[^\s]+(?:\([^\)]*\))?)\s*(?P<rest>.*)$")

def parse_columns(body):
    cols = []
    for line in body.splitlines():
        line = line.strip().rstrip(',')
        if not line.startswith('`'):
            continue
        m = col_re.match(line)
        if not m:
            continue
        col = m.group('col')
        col_type = m.group('type')
        rest = m.group('rest')
        not_null = 'NOT NULL' in rest
        default = ''
        m_def = re.search(r"DEFAULT\s+([^\s,]+)", rest)
        if m_def:
            default = m_def.group(1).strip("'")
        comment = ''
        m_c = re.search(r"COMMENT\s+'([^']*)'", rest)
        if m_c:
            comment = m_c.group(1)
        constraint = []
        constraint.append('NOT NULL' if not_null else 'NULL')
        if default != '':
            constraint.append(f"DEFAULT {default}")
        cols.append({
            'name': col,
            'type': col_type,
            'constraint': ' '.join(constraint).strip(),
            'comment': comment,
        })
    return cols


tables = []
for m in create_re.finditer(sql):
    name = m.group('name')
    body = m.group('body')
    cols = parse_columns(body)
    tables.append((name, cols))

name_to_cols = {n: c for n, c in tables}
ordered_names = [n for n, _ in tables]


def table_desc(name):
    return {
        'ai_conversation': 'AI会话表用于存储AI助手对话记录。',
        'ai_suggestion': 'AI建议表用于存储AI生成的会议智能建议。',
        'app_update': '应用更新表用于存储客户端版本更新信息。',
        'meeting_info': '会议信息表存储会议的详细信息，包括会议号、主题、状态等核心数据。',
        'meeting_member': '会议成员表记录参与会议的用户信息及状态。',
        'meeting_record': '会议记录表存储会议过程与统计信息。',
        'meeting_reserve': '预约会议表存储用户预约的会议信息。',
        'meeting_reserve_member': '预约成员表记录预约会议的参与人员。',
        'meeting_summary': '会议摘要表存储会议摘要与关键要点。',
        'user_contact': '联系人表存储用户的好友关系。',
        'user_contact_apply': '联系人申请表存储好友申请记录。',
        'user_info': '用户信息表用于存储用户的基本账号信息，是系统的基础数据表之一。',
        'user_notification': '通知表存储系统的各类通知信息。',
        'user_notification_backup': '通知备份表用于存储通知历史备份数据。',
        'user_settings': '用户设置表存储用户的个人偏好配置。',
    }[name]


def set_cell_text(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def apply_three_line_table(table):
    tbl = table._element
    tblPr = tbl.tblPr
    # Clear existing borders by setting to nil
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        tblBorders.append(elem)
    tblPr.append(tblBorders)

    def set_row_border(row, top=False, bottom=False):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trBorders = OxmlElement('w:trBorders')
        for edge, enable in (('top', top), ('bottom', bottom)):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single' if enable else 'nil')
            elem.set(qn('w:sz'), '8')
            elem.set(qn('w:color'), '000000')
            trBorders.append(elem)
        trPr.append(trBorders)

    if table.rows:
        set_row_border(table.rows[0], top=True, bottom=True)
        set_row_border(table.rows[-1], bottom=True)


doc = docx.Document(str(doc_path))

body = doc._body._body

# Find start and end element indexes in body list
start_idx = None
end_idx = None

for p in doc.paragraphs:
    if p.text.strip() == '3.3.2 主要数据表设计':
        start_idx = body.index(p._element)
        break

if start_idx is None:
    raise SystemExit('Section start not found')

for p in doc.paragraphs:
    if body.index(p._element) > start_idx and p.text.strip() == '数据库ER图设计':
        end_idx = body.index(p._element)
        break

if end_idx is None:
    raise SystemExit('Section end not found')

# Remove everything between start and end (exclusive)
for i in range(end_idx - 1, start_idx, -1):
    body.remove(body[i])

# Insert new content before end_idx
insert_idx = start_idx + 1

for idx, name in enumerate(ordered_names, start=1):
    desc = table_desc(name)
    ref = f"如表3-{idx}所示。"
    para_text = f"（{idx}）{name}：{desc}{ref}"

    # Paragraph
    p = doc.add_paragraph(para_text)
    body.remove(p._element)
    body.insert(insert_idx, p._element)
    insert_idx += 1

    # Caption paragraph
    cap = doc.add_paragraph('')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = cap.add_run('表')
    run2 = cap.add_run(f"3-{idx}")
    run3 = cap.add_run(f" {name}表")
    for r in (run1, run2, run3):
        r.bold = True
        r.font.size = Pt(10.5)
    run1.font.name = '宋体'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.name = 'Times New Roman'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run3.font.name = '宋体'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    body.remove(cap._element)
    body.insert(insert_idx, cap._element)
    insert_idx += 1

    cols = name_to_cols.get(name, [])
    table = doc.add_table(rows=1 + len(cols), cols=4)
    # Insert table element
    body.remove(table._tbl)
    body.insert(insert_idx, table._tbl)
    insert_idx += 1

    headers = ['字段名', '类型', '约束默认', '说明']
    for cell, h in zip(table.rows[0].cells, headers):
        set_cell_text(cell, h)

    for r, col in enumerate(cols, start=1):
        row = table.rows[r].cells
        set_cell_text(row[0], col['name'])
        set_cell_text(row[1], col['type'])
        set_cell_text(row[2], col['constraint'])
        set_cell_text(row[3], col['comment'])

    apply_three_line_table(table)

# Save
_doc = doc
_doc.save(str(doc_path))
print('updated', doc_path)
