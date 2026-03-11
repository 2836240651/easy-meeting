from pathlib import Path
import docx

in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线_三线表修正.docx")
out_path = in_path.with_name(in_path.stem + '_第6章更新.docx')

doc = docx.Document(str(in_path))
body = doc._body._body

start_idx = None
for p in doc.paragraphs:
    if p.text.strip() == '6.1 工作总结':
        start_idx = body.index(p._element)
        break
if start_idx is None:
    raise SystemExit('6.1 not found')

# remove from start_idx+1 to end
for i in range(len(body)-1, start_idx, -1):
    body.remove(body[i])

insert_idx = start_idx + 1

lines = [
    '本文围绕“基于SpringBoot的多人在线会议系统”完成了系统设计与实现。后端采用Spring Boot与MyBatis进行业务开发，数据层使用MySQL，缓存与会话相关数据由Redis支撑；文件与头像上传集成MinIO。前端采用Vue 3构建界面，通过RESTful API完成业务交互，并结合WebRTC与WebSocket实现实时音视频与消息能力。',
    '系统功能层面完成了用户注册登录、个人信息与设置管理、联系人与申请管理、会议创建与加入、预约会议、会议成员管理、会中聊天与通知等核心模块；同时引入AI助手接口，支持对话、摘要与建议功能，提升会议效率。接口设计按模块划分，规范统一，满足实际项目的扩展与维护需求。',
    '系统测试覆盖用户、会议、音视频、消息、联系人与通知等模块，结果表明系统功能完整、流程可用，能够满足中小规模团队日常会议场景的实际需求。通过项目开发，进一步掌握了WebRTC、多端协同、实时通信与工程化开发流程，也积累了面向实际需求进行架构与模块化设计的经验。',
    '6.2 存在的不足',
    '尽管系统达到预期目标，但仍存在一些不足：',
    '1. 会议规模受限。当前架构更适合中小规模会议，参与人数过多会带来终端负载与带宽压力。',
    '2. 弱网适配能力有限。丢包与抖动场景下音视频体验仍需进一步优化。',
    '3. 终端适配不够完善。移动端仅能通过浏览器访问，功能与体验尚不充分。',
    '4. 会议录制与回放能力不足，难以满足培训、复盘等场景需求。',
    '5. 监控与运维体系较弱，缺少完善的性能指标与告警机制。',
    '6. 安全与合规能力需持续增强，如细粒度权限控制与审计能力。',
    '6.3 未来展望',
    '后续工作将围绕实际使用场景持续迭代：',
    '1. 架构升级。引入SFU/转发服务，提升多方会议的并发规模与稳定性。',
    '2. 网络自适应优化。完善码率控制与丢包恢复策略，提高弱网环境下的可用性。',
    '3. 多端能力增强。补齐移动端功能，优化小屏交互与推送体验。',
    '4. 会议沉淀能力扩展。支持录制、回放、检索与会议资料管理。',
    '5. 智能化能力增强。结合现有AI接口拓展会议纪要、要点提取、行动项跟踪等功能。',
    '6. 运维体系完善。引入统一监控、日志追踪与报警机制，提升系统可维护性。',
]

for text in lines:
    p = doc.add_paragraph(text)
    body.remove(p._element)
    body.insert(insert_idx, p._element)
    insert_idx += 1

_doc = doc
_doc.save(str(out_path))
print('updated', out_path)
