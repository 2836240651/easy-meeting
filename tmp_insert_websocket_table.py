from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path(r"C:\Users\13377\Desktop\论文相关\第2章_相关技术介绍_已调整格式_v2.docx")
OUT = Path(r"C:\Users\13377\Desktop\论文相关\第2章_相关技术介绍_已调整格式_v3.docx")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_run(run, east_asia, size_pt, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def style_paragraph(paragraph, center=False):
    fmt = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = None if center else Cm(0.74)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5


doc = Document(str(SRC))
anchor = None
for para in doc.paragraphs:
    if para.text.strip() == "WebSocket和HTTP是两种不同的通信协议，主要区别如下：":
        anchor = para
        break

if anchor is None:
    raise SystemExit("Anchor paragraph not found")

title_para = deepcopy(anchor._p)
for child in list(title_para):
    title_para.remove(child)
anchor._p.addnext(title_para)

title_p = None
for p in doc.paragraphs:
    if p._p is title_para:
        title_p = p
        break
if title_p is None:
    raise SystemExit("Title paragraph wrapper not found")

style_paragraph(title_p, center=True)
run = title_p.add_run("表2-1 WebSocket与HTTP区别")
set_run(run, "宋体", 12, False)

table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ["比较维度", "HTTP", "WebSocket"]
rows = [
    ["连接方式", "请求-响应，通常一次请求对应一次连接或短连接交互。", "建立连接后保持长连接，支持持续双向通信。"],
    ["通信方向", "以客户端主动发起请求为主，服务端被动响应。", "客户端和服务端都可以主动发送消息。"],
    ["实时性", "实时性较弱，通常需要轮询或长轮询实现准实时通信。", "实时性强，适合即时消息和信令交互场景。"],
    ["数据开销", "每次请求都携带较完整的 HTTP 头部，开销相对较大。", "建立连接后帧头较小，持续通信时开销更低。"],
    ["适用场景", "适合页面访问、资源获取、普通业务接口调用。", "适合在线聊天、视频会议、实时通知等高实时场景。"],
]

widths = [Cm(3.2), Cm(6.4), Cm(6.4)]
for col, width in zip(table.columns, widths):
    for cell in col.cells:
        cell.width = width

for j, text in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, "黑体", 12, False)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i, row_data in enumerate(rows, start=1):
    for j, text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.first_line_indent = None
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.25
        r = p.add_run(text)
        set_run(r, "宋体", 12, False)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for row in table.rows:
    for cell in row.cells:
        set_cell_border(
            cell,
            left={"val": "nil"},
            right={"val": "nil"},
        )

for cell in table.rows[0].cells:
    set_cell_border(
        cell,
        top={"val": "single", "sz": 12, "color": "000000"},
        bottom={"val": "single", "sz": 8, "color": "000000"},
        left={"val": "nil"},
        right={"val": "nil"},
    )

for cell in table.rows[-1].cells:
    set_cell_border(
        cell,
        bottom={"val": "single", "sz": 12, "color": "000000"},
        left={"val": "nil"},
        right={"val": "nil"},
    )

anchor._p.addnext(table._tbl)
title_p._p.addnext(table._tbl)

doc.save(str(OUT))
print(OUT)
