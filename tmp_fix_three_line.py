from pathlib import Path
import re
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新_接口设计按代码更新_5_2_1表格三线.docx")
if not in_path.exists():
    raise SystemExit('input doc not found')

out_path = in_path.with_name(in_path.stem + '_三线表修正.docx')

doc = docx.Document(str(in_path))
body = doc._body._body

# locate section bounds
start_idx = None
end_idx = None
for p in doc.paragraphs:
    if p.text.strip() == '5.2.1 用户模块测试':
        start_idx = body.index(p._element)
        break
if start_idx is None:
    raise SystemExit('5.2.1 not found')
for p in doc.paragraphs:
    if body.index(p._element) > start_idx and p.text.strip().startswith('5.2.2'):
        end_idx = body.index(p._element)
        break
if end_idx is None:
    raise SystemExit('5.2.2 not found')

# helper

def clear_tbl_borders(tbl):
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        tblBorders.append(elem)
    tblPr.append(tblBorders)


def set_row_border(row, top=False, bottom=False):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trBorders = OxmlElement('w:trBorders')
    for edge, enable in (('top', top), ('bottom', bottom)):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single' if enable else 'nil')
        elem.set(qn('w:sz'), '8')
        elem.set(qn('w:color'), '000000')
        trBorders.append(elem)
    trPr.append(trBorders)


def set_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    noWrap = OxmlElement('w:noWrap')
    tcPr.append(noWrap)


def truncate_text(text, max_len=25):
    t = text.strip()
    if len(t) <= max_len:
        return t
    return t[:max_len] + '…'

# find tables in section
section_tbl_idxs = []
for idx in range(start_idx + 1, end_idx):
    elem = body[idx]
    if elem.tag.endswith('}tbl'):
        section_tbl_idxs.append(idx)

# update each table
for tbl_idx in section_tbl_idxs:
    table = None
    for t in doc.tables:
        if t._tbl is body[tbl_idx]:
            table = t
            break
    if table is None:
        continue

    clear_tbl_borders(table._element)

    # apply three-line: only header top/bottom, last row bottom
    if table.rows:
        set_row_border(table.rows[0], top=True, bottom=True)
        set_row_border(table.rows[-1], bottom=True)

    # enforce no wrap and truncate long text
    for row in table.rows:
        for cell in row.cells:
            set_no_wrap(cell)
            # truncate each paragraph text in cell
            for p in cell.paragraphs:
                if p.text:
                    p.text = truncate_text(p.text, max_len=25)

# save
_doc = doc
_doc.save(str(out_path))
print('updated', out_path)
