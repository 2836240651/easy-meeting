from pathlib import Path
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SONGTI = "\u5b8b\u4f53"
HEITI = "\u9ed1\u4f53"
DENGXIAN = "\u7b49\u7ebf"
SUMMARY_TITLE = "\u672c\u7ae0\u5c0f\u7ed3"
REFERENCES_TITLE = "\u53c2\u8003\u6587\u732e"


def ensure_rpr(style):
    style_el = style._element
    rpr = style_el.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style_el.append(rpr)
    return rpr


def set_style_font(style, ascii_font: str, east_asia_font: str, size_pt: float, bold: bool = False):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0)
    rpr = ensure_rpr(style)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_run_font(run, east_asia_font: str, size_pt: float, bold: bool = False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def append_page_number(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)
    run.font.name = SONGTI
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), SONGTI)
    rfonts.set(qn("w:ascii"), SONGTI)
    rfonts.set(qn("w:hAnsi"), SONGTI)


def format_paragraph(paragraph, style_name: str):
    fmt = paragraph.paragraph_format
    fmt.left_indent = None
    fmt.right_indent = None

    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = None
        fmt.hanging_indent = None
        fmt.space_before = Pt(18)
        fmt.space_after = Pt(18)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif style_name == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.hanging_indent = None
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
        fmt.line_spacing = 1.5
    elif style_name == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.hanging_indent = None
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(3)
        fmt.line_spacing = 1.5
    elif style_name == "Source Code":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = None
        fmt.hanging_indent = None
        fmt.left_indent = Cm(0.74)
        fmt.right_indent = Cm(0.74)
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(3)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif style_name == "Reference":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = None
        fmt.hanging_indent = Cm(0.74)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(0.74)
        fmt.hanging_indent = None
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5


def apply_styles(doc: Document):
    styles = doc.styles
    set_style_font(styles["Normal"], "Times New Roman", SONGTI, 12)
    set_style_font(styles["Body Text"], "Times New Roman", SONGTI, 12)
    set_style_font(styles["First Paragraph"], "Times New Roman", SONGTI, 12)
    set_style_font(styles["Heading 1"], "Times New Roman", HEITI, 15, True)
    set_style_font(styles["Heading 2"], "Times New Roman", HEITI, 14, False)
    set_style_font(styles["Heading 3"], "Times New Roman", HEITI, 12, False)
    if "Source Code" in styles:
        set_style_font(styles["Source Code"], "Consolas", DENGXIAN, 10.5)


def apply_page_layout(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragraph.paragraph_format
        fmt.first_line_indent = None
        fmt.hanging_indent = None
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        append_page_number(paragraph)


def format_document(path: Path):
    backup_path = path.with_name(path.stem + "_format_backup.docx")
    output_path = path.with_name(path.stem + "_已调整格式.docx")

    if not backup_path.exists():
        shutil.copy2(path, backup_path)

    doc = Document(str(path))
    apply_styles(doc)
    apply_page_layout(doc)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name
        if text == REFERENCES_TITLE:
            paragraph.style = doc.styles["Heading 1"]
            style_name = "Heading 1"
        elif text == SUMMARY_TITLE:
            paragraph.style = doc.styles["Heading 2"]
            style_name = "Heading 2"
        elif text.startswith("[") and style_name in {"Body Text", "First Paragraph", "Normal"}:
            style_name = "Reference"
        elif style_name not in {"Heading 1", "Heading 2", "Heading 3", "Source Code"}:
            style_name = "Body"

        format_paragraph(paragraph, style_name)

        for run in paragraph.runs:
            if style_name == "Heading 1":
                set_run_font(run, HEITI, 15, True)
            elif style_name == "Heading 2":
                set_run_font(run, HEITI, 14, False)
            elif style_name == "Heading 3":
                set_run_font(run, HEITI, 12, False)
            elif style_name == "Source Code":
                set_run_font(run, DENGXIAN, 10.5, False)
            else:
                set_run_font(run, SONGTI, 12, False)

    try:
        doc.save(str(output_path))
    except PermissionError:
        output_path = path.with_name(path.stem + "_已调整格式_v2.docx")
        doc.save(str(output_path))
    print(f"Formatted: {output_path}")
    print(f"Backup: {backup_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit("Usage: python tmp_docx_format_generic.py <docx-path> [<docx-path> ...]")

    for arg in sys.argv[1:]:
        format_document(Path(arg))
