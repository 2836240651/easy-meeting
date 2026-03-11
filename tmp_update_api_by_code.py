import re, glob, os
from pathlib import Path
import docx

root = r'D:\\JavaPartical\\easymeeting-java\\src\\main\\java\\com\\easymeeting'
controllers = glob.glob(root + r'\\**\\*Controller*.java', recursive=True)

# Helpers to extract mappings
str_lit = r'"([^"\\]*(?:\\.[^"\\]*)*)"'

array_pattern = re.compile(r'\{\s*(' + str_lit + r'(?:\s*,\s*' + str_lit + r')*)\s*\}')
string_pattern = re.compile(str_lit)


def extract_paths(annotation):
    # Try array of strings first
    m = array_pattern.search(annotation)
    if m:
        return re.findall(str_lit, m.group(1))
    # Try value= or direct string
    m = re.search(r'value\s*=\s*' + str_lit, annotation)
    if m:
        return [m.group(1)]
    m = string_pattern.search(annotation)
    if m:
        return [m.group(1)]
    return ['']


def extract_methods(annotation):
    # For Get/Post/Put/DeleteMapping
    if annotation.startswith('@GetMapping'):
        return ['GET']
    if annotation.startswith('@PostMapping'):
        return ['POST']
    if annotation.startswith('@PutMapping'):
        return ['PUT']
    if annotation.startswith('@DeleteMapping'):
        return ['DELETE']
    # For RequestMapping with method=
    m = re.search(r'method\s*=\s*\{?([^}]+)\}?', annotation)
    if m:
        parts = m.group(1)
        methods = re.findall(r'RequestMethod\.([A-Z]+)', parts)
        return methods if methods else ['ALL']
    return ['ALL']


def join_paths(base, sub):
    if not base:
        base = ''
    if not sub:
        return base or '/'
    if base.endswith('/') and sub.startswith('/'):
        return base[:-1] + sub
    if (not base.endswith('/')) and (not sub.startswith('/')):
        return base + '/' + sub
    return base + sub

endpoints = []

for f in controllers:
    txt = Path(f).read_text(encoding='utf-8', errors='ignore')
    base = ''
    # class-level RequestMapping
    m = re.search(r'@RequestMapping\(([^)]*)\)', txt)
    if m:
        paths = extract_paths(m.group(0))
        base = paths[0] if paths else ''

    # find method-level mappings
    # naive: iterate lines and capture annotations before method
    lines = txt.splitlines()
    ann_buf = []
    for line in lines:
        line = line.strip()
        if line.startswith('@'):
            ann_buf.append(line)
            continue
        # method signature heuristic
        if re.match(r'(public|private|protected)\s+[\w<>\[\]]+\s+\w+\s*\(', line):
            # process annotations in buffer
            for ann in ann_buf:
                if ann.startswith('@RequestMapping') or ann.startswith('@GetMapping') or ann.startswith('@PostMapping') or ann.startswith('@PutMapping') or ann.startswith('@DeleteMapping'):
                    paths = extract_paths(ann)
                    methods = extract_methods(ann)
                    for p in paths:
                        full = join_paths(base, p)
                        endpoints.append((full, methods))
            ann_buf = []
        elif line == '':
            # keep annotations across blank lines
            continue
        else:
            # reset buffer if other code appears
            if ann_buf and not line.startswith('@'):
                # keep until method signature
                pass

# normalize endpoints
normalized = {}
for path, methods in endpoints:
    # normalize double slashes
    path = re.sub(r'//+', '/', path)
    if not path.startswith('/'):
        path = '/' + path
    key = path
    cur = normalized.get(key, set())
    for m in methods:
        cur.add(m)
    normalized[key] = cur

# Group by module based on path prefix
module_rules = [
    ('用户模块主要接口：', ['/account', '/userInfo', '/api/settings']),
    ('会议模块主要接口：', ['/meetingInfo', '/meetingMember']),
    ('预约模块主要接口：', ['/meetingReserve', '/meetingReserveMember']),
    ('联系人模块主要接口：', ['/userContact']),
    ('通知模块主要接口：', ['/notification']),
    ('AI模块主要接口：', ['/ai']),
    ('聊天模块主要接口：', ['/chat', '/meetingChatMessage']),
    ('文件与头像接口：', ['/upload', '/files', '/file', '/fileAccess']),
    ('版本更新接口：', ['/admin']),
]

modules = {title: [] for title, _ in module_rules}
others = []

for path, methods in sorted(normalized.items()):
    added = False
    for title, prefixes in module_rules:
        if any(path.startswith(p) for p in prefixes):
            modules[title].append((path, methods))
            added = True
            break
    if not added:
        others.append((path, methods))

# Build lines for doc
lines = []
for title, _ in module_rules:
    items = modules[title]
    if not items:
        continue
    lines.append(title)
    for path, methods in items:
        method_str = '、'.join(sorted(methods))
        lines.append(f"{path}（{method_str}）")
    lines.append('')

if others:
    lines.append('其他接口：')
    for path, methods in others:
        method_str = '、'.join(sorted(methods))
        lines.append(f"{path}（{method_str}）")
    lines.append('')

# Update doc
in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7_接口设计更新.docx")
if not in_path.exists():
    in_path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")

out_path = in_path.with_name(in_path.stem + '_接口设计按代码更新.docx')

_doc = docx.Document(str(in_path))
body = _doc._body._body

start_idx = None
end_idx = None
for p in _doc.paragraphs:
    if p.text.strip() == '3.4.2 主要接口列表':
        start_idx = body.index(p._element)
        break
if start_idx is None:
    raise SystemExit('3.4.2 not found')

for p in _doc.paragraphs:
    if body.index(p._element) > start_idx and p.text.strip() == '本章小结':
        end_idx = body.index(p._element)
        break
if end_idx is None:
    raise SystemExit('end not found')

for i in range(end_idx - 1, start_idx, -1):
    body.remove(body[i])

insert_idx = start_idx + 1
for text in lines:
    p = _doc.add_paragraph(text)
    body.remove(p._element)
    body.insert(insert_idx, p._element)
    insert_idx += 1

_doc.save(str(out_path))
print('updated', out_path)
