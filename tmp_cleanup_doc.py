import re
from pathlib import Path
import docx

path = Path(r"C:\Users\13377\Desktop\论文相关\毕业论back_备份_修正扩展名_论文组织结构已补全_条标题正文统一格式_v7.docx")

doc = docx.Document(str(path))
body = doc._body._body

# Find ER heading paragraph element index
start = None
for p in doc.paragraphs:
    if p.text.strip() == '数据库ER图设计':
        start = body.index(p._element)
        break

if start is None:
    raise SystemExit('ER heading not found')

# Find end marker paragraph index (the last relationship line) or next section heading if any
end = None
for p in doc.paragraphs:
    if body.index(p._element) > start and p.text.strip().startswith('会议与消息：'):
        end = body.index(p._element)
        break

if end is None:
    end = len(body) - 1

# Remove duplicated table entries between ER heading and end marker
remove_idxs = []
for p in doc.paragraphs:
    idx = body.index(p._element)
    if idx <= start or idx >= end:
        continue
    t = p.text.strip()
    if re.match(r'^（\d+）', t) or t.startswith('表3-'):
        remove_idxs.append(idx)

for idx in sorted(set(remove_idxs), reverse=True):
    body.remove(body[idx])

# Save
_doc = doc
_doc.save(str(path))
print('cleaned', path)
